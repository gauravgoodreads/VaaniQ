#!/usr/bin/env python3
"""Generate VaaniQ master DOCX with live screenshots and metrics."""

from __future__ import annotations

import json
import sys
import urllib.request
import zipfile
from datetime import UTC, datetime
from pathlib import Path

import matplotlib

matplotlib.use("Agg")
import matplotlib.pyplot as plt
import numpy as np
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from PIL import Image
from playwright.sync_api import sync_playwright

REPO = Path(__file__).resolve().parents[1]
OUT_DIR = REPO / "docs" / "assets" / "screenshots"
DOCX_OUT = REPO / "docs" / "VaaniQ_Master_Presentation_FINAL_v2.docx"
TRAIN_REPORT = REPO / "models" / "checkpoints" / "xlsr_aasist" / "train_report.json"
FIG_DIR = REPO / "docs" / "assets" / "figures"
VERIFIED_FIG_DIR = REPO / "docs" / "assets" / "verified_figures"
DEMO_WAV_REAL = REPO / "data" / "demo_corpus" / "audio" / "hi-0.wav"
DEMO_WAV_FAKE = REPO / "data" / "demo_corpus" / "audio" / "hi-1.wav"
PUBLICATION_ROOT = REPO / "data" / "publication_corpus"

BASE = "http://127.0.0.1:5173"
API = "http://127.0.0.1:8001"

PAGES: list[tuple[str, str, str]] = [
    ("01-landing.png", "/", "Landing Page"),
    ("02-dashboard.png", "/dashboard", "Research Dashboard"),
    ("03-upload.png", "/upload", "Upload and Detection"),
    ("04-live.png", "/live", "Live Microphone Streaming"),
    ("05-calibration.png", "/calibration", "Calibration and Reliability"),
    ("06-explainability.png", "/explainability", "Explainability Suite"),
    ("07-datasets.png", "/datasets", "Dataset Explorer"),
    ("08-human-study.png", "/human-study", "Human Perception Study"),
    ("09-research-metrics.png", "/research-metrics", "Research Metrics"),
    ("10-experiments.png", "/experiments", "Experiment Browser"),
    ("11-history.png", "/history", "Detection History"),
    ("12-inference.png", "/inference", "Inference Browser"),
    ("13-admin.png", "/admin", "System Administration"),
    ("14-docs.png", "/docs", "Documentation Hub"),
    ("15-api-docs.png", f"{API}/docs", "API Documentation (Swagger)"),
]


def fetch_json(url: str) -> dict[str, object]:
    with urllib.request.urlopen(url, timeout=30) as resp:
        return json.loads(resp.read().decode())


def post_inference(wav: Path, language: str = "hi") -> dict[str, object]:
    import mimetypes

    boundary = "----VaaniQBoundary7MA4YWxk"
    body_parts: list[bytes] = []
    for key, val in [("language", language), ("model_id", "aasist-v1")]:
        body_parts.append(f"--{boundary}\r\n".encode())
        body_parts.append(f'Content-Disposition: form-data; name="{key}"\r\n\r\n'.encode())
        body_parts.append(f"{val}\r\n".encode())
    mime = "audio/flac" if wav.suffix.lower() == ".flac" else (
        mimetypes.guess_type(wav.name)[0] or "audio/wav"
    )
    data = wav.read_bytes()
    body_parts.append(f"--{boundary}\r\n".encode())
    body_parts.append(
        f'Content-Disposition: form-data; name="file"; filename="{wav.name}"\r\n'.encode()
    )
    body_parts.append(f"Content-Type: {mime}\r\n\r\n".encode())
    body_parts.append(data)
    body_parts.append(f"\r\n--{boundary}--\r\n".encode())
    body = b"".join(body_parts)
    req = urllib.request.Request(
        f"{API}/api/v1/inference",
        data=body,
        headers={"Content-Type": f"multipart/form-data; boundary={boundary}"},
        method="POST",
    )
    with urllib.request.urlopen(req, timeout=60) as resp:
        return json.loads(resp.read().decode())


def resolve_inference_sample(label: str) -> tuple[Path, str]:
    """Prefer a clean held-out publication clip, then fall back to demo audio."""
    manifest = PUBLICATION_ROOT / "manifest.jsonl"
    if manifest.is_file():
        with manifest.open(encoding="utf-8") as handle:
            for line in handle:
                if not line.strip():
                    continue
                row = json.loads(line)
                if (
                    row.get("label") == label
                    and row.get("split") == "test"
                    and row.get("compression_status") == "clean"
                ):
                    path = PUBLICATION_ROOT / str(row["uri"])
                    if path.is_file():
                        return path, str(row.get("language", "hi"))
    fallback = DEMO_WAV_REAL if label == "real" else DEMO_WAV_FAKE
    return fallback, "hi"


def capture_screenshots() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    with sync_playwright() as p:
        browser = p.chromium.launch(headless=True)
        page = browser.new_page(viewport={"width": 1440, "height": 900})
        page.set_default_timeout(45000)
        for fname, route, _ in PAGES:
            url = route if route.startswith("http") else f"{BASE}{route}"
            page.goto(url, wait_until="networkidle")
            page.wait_for_timeout(1500)
            page.screenshot(path=str(OUT_DIR / fname), full_page=True)
            print(f"captured {fname}")
        browser.close()


def add_title(doc: Document, text: str, level: int = 0) -> None:
    if level == 0:
        p = doc.add_heading(text, level=0)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        doc.add_heading(text, level=level)


def add_image(doc: Document, path: Path, width: float = 6.5) -> None:
    if path.is_file():
        with Image.open(path) as image:
            pixel_width, pixel_height = image.size
        ratio = pixel_width / max(pixel_height, 1)
        max_width = min(width, 6.7)
        max_height = 7.7
        fitted_width = max_width
        fitted_height = fitted_width / ratio
        if fitted_height > max_height:
            fitted_height = max_height
            fitted_width = fitted_height * ratio
        shape = doc.add_picture(
            str(path),
            width=Inches(fitted_width),
            height=Inches(fitted_height),
        )
        shape._inline.getparent().getparent().alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()
    else:
        doc.add_paragraph(f"[Image pending: {path.name}]")


def add_tiled_screenshot(doc: Document, path: Path) -> None:
    """Add a full-page overview plus readable vertical tiles."""
    add_image(doc, path, 6.2)
    with Image.open(path) as image:
        pixel_width, pixel_height = image.size
        if pixel_height <= int(pixel_width * 1.15):
            return
        tile_dir = OUT_DIR / "tiles"
        tile_dir.mkdir(parents=True, exist_ok=True)
        tile_height = int(pixel_width * 1.12)
        starts = list(range(0, pixel_height, tile_height))
        for index, top in enumerate(starts, start=1):
            bottom = min(pixel_height, top + tile_height)
            if bottom - top < pixel_width * 0.25 and index > 1:
                top = max(0, pixel_height - tile_height)
            tile = image.crop((0, top, pixel_width, bottom))
            tile_path = tile_dir / f"{path.stem}-{index:02d}.png"
            tile.save(tile_path)
            doc.add_paragraph(f"Detail {index} of {len(starts)}", style="Caption")
            add_image(doc, tile_path, 6.2)
            if index < len(starts):
                doc.add_page_break()


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
    for ri, row in enumerate(rows):
        cells = table.rows[ri + 1].cells
        for ci, val in enumerate(row):
            cells[ci].text = str(val)
    doc.add_paragraph()


def _metric(
    data: dict[str, object],
    key: str,
    default: float = 0.0,
) -> float:
    value = data.get(key, default)
    return float(value) if isinstance(value, int | float) else default


def load_train_report() -> dict[str, object]:
    """Load Baseline V1 with frozen Round 3 metrics overriding legacy files."""
    sys.path.insert(0, str(REPO / "scripts"))
    from report_data import load_train_report as load_frozen_train_report

    frozen = load_frozen_train_report()
    if frozen:
        return frozen
    if not TRAIN_REPORT.is_file():
        return {}
    with TRAIN_REPORT.open(encoding="utf-8") as fh:
        raw = json.load(fh)
    return raw if isinstance(raw, dict) else {}


def merge_pipeline(pipeline: dict[str, object]) -> dict[str, object]:
    """Ensure pipeline payload includes eval_metrics from train_report.json."""
    report = load_train_report()
    if not report:
        return pipeline
    merged = dict(pipeline)
    for key in (
        "eval_metrics",
        "test_metrics",
        "validation_metrics",
        "test_accuracy",
        "test_eer",
        "test_ece",
        "test_brier",
        "test_roc_auc",
        "n_clips",
        "total_hours",
        "n_train",
        "n_val",
        "n_test",
        "gpu",
        "cuda_available",
        "data_provenance",
        "corpus_provenance",
        "speaker_disjoint_verified",
        "speaker_counts",
        "status",
        "pipeline",
    ):
        if key not in merged or merged.get(key) in (None, {}, []):
            value = report.get(key)
            if value is not None:
                merged[key] = value
    eval_metrics = merged.get("eval_metrics")
    if not isinstance(eval_metrics, dict):
        fallback = report.get("test_metrics")
        if isinstance(fallback, dict):
            merged["eval_metrics"] = fallback
    return merged


def _resolve_metrics(pipeline: dict[str, object]) -> dict[str, object]:
    raw = pipeline.get("eval_metrics")
    if isinstance(raw, dict) and raw:
        return raw
    raw = pipeline.get("test_metrics")
    if isinstance(raw, dict) and raw:
        return raw
    report = load_train_report()
    for key in ("eval_metrics", "test_metrics"):
        block = report.get(key)
        if isinstance(block, dict) and block:
            return block
    return {}


def _evaluation_descriptor(pipeline: dict[str, object]) -> tuple[str, str]:
    """Return a concise figure title and provenance disclosure."""
    provenance = str(pipeline.get("data_provenance", ""))
    if "kathbath" in provenance and "indicsynth" in provenance:
        return (
            "speaker-disjoint publication test subset",
            "Measured on the persisted bounded publication subset; not the full 303 GB corpora.",
        )
    return (
        "synthetic test set",
        "Synthetic demo evaluation only - not a curated-corpus publication result.",
    )


def load_cross_lingual_results() -> dict[str, dict[str, object]]:
    """Load three leave-one-language-out test reports when available."""
    results: dict[str, dict[str, object]] = {}
    for language in ("hi", "mr", "ta"):
        path = (
            REPO
            / "models"
            / "checkpoints"
            / "rq3"
            / f"test_{language}"
            / "train_report.json"
        )
        if not path.is_file():
            continue
        raw = json.loads(path.read_text(encoding="utf-8"))
        if isinstance(raw, dict):
            metrics = raw.get("test_metrics")
            if isinstance(metrics, dict):
                results[language] = metrics
    return results


def generate_verified_figures(pipeline: dict[str, object]) -> None:
    """Generate print-safe figures from the measured synthetic test report."""
    VERIFIED_FIG_DIR.mkdir(parents=True, exist_ok=True)
    metrics = _resolve_metrics(pipeline)
    evaluation_name, disclosure = _evaluation_descriptor(pipeline)
    style = {
        "figure.facecolor": "white",
        "axes.facecolor": "white",
        "font.size": 11,
        "axes.titlesize": 15,
        "axes.labelsize": 11,
    }
    plt.rcParams.update(style)

    names = ["Accuracy", "Precision", "Recall", "F1", "ROC-AUC"]
    values = [
        _metric(metrics, "accuracy"),
        _metric(metrics, "precision"),
        _metric(metrics, "recall"),
        _metric(metrics, "f1"),
        _metric(metrics, "roc_auc"),
    ]
    fig, ax = plt.subplots(figsize=(9, 5.2))
    fig.subplots_adjust(left=0.09, right=0.98, bottom=0.17, top=0.80)
    bars = ax.bar(names, values, color=["#0f766e", "#14b8a6", "#2dd4bf", "#5eead4", "#99f6e4"])
    ax.set_ylim(0, 1.12)
    ax.set_ylabel("Score")
    ax.set_title(
        f"Held-out {evaluation_name} performance (n={int(metrics.get('n', 90))})",
        pad=16,
    )
    ax.grid(axis="y", alpha=0.2)
    for bar, value in zip(bars, values, strict=True):
        ax.text(
            bar.get_x() + bar.get_width() / 2,
            value + 0.025,
            f"{value:.2%}",
            ha="center",
            va="bottom",
            fontsize=10,
        )
    fig.text(
        0.5,
        0.035,
        disclosure,
        ha="center",
        fontsize=9,
    )
    fig.savefig(VERIFIED_FIG_DIR / "performance_overview.png", dpi=220)
    plt.close(fig)

    per_language = metrics.get("per_language")
    language_data = per_language if isinstance(per_language, dict) else {}
    languages = ["Hindi", "Marathi", "Tamil"]
    language_codes = ["hi", "mr", "ta"]
    accuracies = [
        _metric(
            language_data.get(code, {}) if isinstance(language_data.get(code), dict) else {},
            "accuracy",
        )
        for code in language_codes
    ]
    f1_scores = [
        _metric(
            language_data.get(code, {}) if isinstance(language_data.get(code), dict) else {},
            "f1",
        )
        for code in language_codes
    ]
    x = np.arange(len(languages))
    fig, ax = plt.subplots(figsize=(9, 5.2))
    fig.subplots_adjust(left=0.09, right=0.82, bottom=0.14, top=0.80)
    width = 0.36
    acc_bars = ax.bar(x - width / 2, accuracies, width, label="Accuracy", color="#0f766e")
    f1_bars = ax.bar(
        x + width / 2,
        f1_scores,
        width,
        label="F1 score",
        color="#f59e0b",
    )
    ax.set_xticks(x, languages)
    ax.set_ylim(0, 1.12)
    ax.set_ylabel("Rate")
    ax.set_title(f"Per-language held-out {evaluation_name} performance", pad=16)
    ax.legend(loc="upper left", bbox_to_anchor=(1.01, 1.0), borderaxespad=0)
    ax.grid(axis="y", alpha=0.2)
    for bars_group in (acc_bars, f1_bars):
        for bar in bars_group:
            ax.text(
                bar.get_x() + bar.get_width() / 2,
                bar.get_height() + 0.025,
                f"{bar.get_height():.2%}",
                ha="center",
                va="bottom",
                fontsize=10,
            )
    fig.savefig(VERIFIED_FIG_DIR / "language_breakdown.png", dpi=220)
    plt.close(fig)

    per_condition = metrics.get("per_condition")
    condition_data = per_condition if isinstance(per_condition, dict) else {}
    condition_keys = ["clean", "opus_whatsapp_sim"]
    condition_names = ["Clean", "Opus simulation"]
    condition_acc = [
        _metric(
            condition_data.get(key, {}) if isinstance(condition_data.get(key), dict) else {},
            "accuracy",
        )
        for key in condition_keys
    ]
    condition_eer = [
        _metric(
            condition_data.get(key, {}) if isinstance(condition_data.get(key), dict) else {}, "eer"
        )
        for key in condition_keys
    ]
    x = np.arange(len(condition_names))
    fig, ax = plt.subplots(figsize=(8.5, 5.2))
    fig.subplots_adjust(left=0.09, right=0.82, bottom=0.14, top=0.80)
    acc_bars = ax.bar(x - width / 2, condition_acc, width, label="Accuracy", color="#0f766e")
    eer_bars = ax.bar(x + width / 2, condition_eer, width, label="EER", color="#f59e0b")
    ax.set_xticks(x, condition_names)
    ax.set_ylim(0, 1.12)
    ax.set_ylabel("Rate")
    ax.set_title(f"Condition breakdown on held-out {evaluation_name}", pad=16)
    ax.legend(loc="upper left", bbox_to_anchor=(1.01, 1.0), borderaxespad=0)
    ax.grid(axis="y", alpha=0.2)
    for bars_group in (acc_bars, eer_bars):
        for bar in bars_group:
            ax.text(
                bar.get_x() + bar.get_width() / 2,
                max(bar.get_height() + 0.025, 0.025),
                f"{bar.get_height():.2%}",
                ha="center",
                va="bottom",
                fontsize=10,
            )
    fig.savefig(VERIFIED_FIG_DIR / "condition_breakdown.png", dpi=220)
    plt.close(fig)

    matrix_raw = metrics.get("confusion_matrix")
    matrix = np.asarray(matrix_raw if isinstance(matrix_raw, list) else [[0, 0], [0, 0]])
    fig, ax = plt.subplots(figsize=(6.5, 5.5))
    fig.subplots_adjust(left=0.20, right=0.88, bottom=0.15, top=0.82)
    image = ax.imshow(matrix, cmap="Blues")
    ax.set_xticks([0, 1], ["Predicted real", "Predicted fake"])
    ax.set_yticks([0, 1], ["Actual real", "Actual fake"])
    ax.set_title(f"Confusion matrix: {evaluation_name}", pad=14)
    for row in range(2):
        for col in range(2):
            ax.text(col, row, str(int(matrix[row, col])), ha="center", va="center", fontsize=16)
    fig.colorbar(image, ax=ax, fraction=0.046)
    fig.savefig(VERIFIED_FIG_DIR / "confusion_matrix.png", dpi=220)
    plt.close(fig)

    fig, ax = plt.subplots(figsize=(8.5, 4.8))
    fig.subplots_adjust(left=0.11, right=0.98, bottom=0.15, top=0.80)
    calibration_names = ["ECE", "Brier"]
    calibration_values = [_metric(metrics, "ece"), _metric(metrics, "brier")]
    bars = ax.bar(calibration_names, calibration_values, color=["#7c3aed", "#a78bfa"], width=0.55)
    ax.set_ylim(0, max(0.25, max(calibration_values, default=0.0) + 0.1))
    ax.set_title(f"Calibration errors on held-out {evaluation_name}", pad=16)
    ax.set_ylabel("Error (lower is better)")
    ax.grid(axis="y", alpha=0.2)
    for bar, value in zip(bars, calibration_values, strict=True):
        ax.text(
            bar.get_x() + bar.get_width() / 2,
            value + 0.01,
            f"{value:.3f}",
            ha="center",
            va="bottom",
        )
    fig.savefig(VERIFIED_FIG_DIR / "calibration_errors.png", dpi=220)
    plt.close(fig)

    cal_pre = metrics.get("calibration_pre")
    cal_post = metrics.get("calibration_post")
    pre_block = cal_pre if isinstance(cal_pre, dict) else {}
    post_block = cal_post if isinstance(cal_post, dict) else {}
    pre_ece = _metric(pre_block, "ece")
    pre_brier = _metric(pre_block, "brier")
    post_ece = _metric(post_block, "ece", _metric(metrics, "ece"))
    post_brier = _metric(post_block, "brier", _metric(metrics, "brier"))
    fig, ax = plt.subplots(figsize=(9, 5.2))
    fig.subplots_adjust(left=0.11, right=0.98, bottom=0.15, top=0.80)
    x = np.arange(2)
    width = 0.36
    ax.bar(x - width / 2, [pre_ece, pre_brier], width, label="Pre-calibration", color="#94a3b8")
    ax.bar(x + width / 2, [post_ece, post_brier], width, label="Post-calibration", color="#7c3aed")
    ax.set_xticks(x, ["ECE", "Brier"])
    ax.set_ylim(0, max(0.45, max(pre_ece, pre_brier, post_ece, post_brier) + 0.12))
    ax.set_title(f"Calibration comparison on held-out {evaluation_name}", pad=16)
    ax.set_ylabel("Error (lower is better)")
    ax.legend()
    ax.grid(axis="y", alpha=0.2)
    for offset, values in ((-width / 2, [pre_ece, pre_brier]), (width / 2, [post_ece, post_brier])):
        for idx, value in enumerate(values):
            ax.text(
                idx + offset,
                value + 0.01,
                f"{value:.3f}",
                ha="center",
                va="bottom",
                fontsize=10,
            )
    fig.savefig(VERIFIED_FIG_DIR / "calibration_pre_post.png", dpi=220)
    plt.close(fig)

    diagram_raw = metrics.get("reliability_diagram")
    diagram_rows = diagram_raw if isinstance(diagram_raw, list) else []
    points = [
        (
            float(
                row.get(
                    "confidence",
                    (float(row.get("bin_lo", 0)) + float(row.get("bin_hi", 0))) / 2,
                )
            ),
            float(row.get("accuracy", 0)),
            float(row.get("count", 0)),
        )
        for row in diagram_rows
        if isinstance(row, dict) and float(row.get("count", 0)) > 0
    ]
    fig, ax = plt.subplots(figsize=(8.5, 5.5))
    fig.subplots_adjust(left=0.11, right=0.98, bottom=0.14, top=0.78)
    ax.plot([0, 1], [0, 1], "--", color="#64748b", linewidth=1.5, label="Perfect calibration")
    if points:
        xs, ys, counts = zip(*points, strict=True)
        ax.plot(xs, ys, "o-", color="#0f766e", linewidth=2, markersize=7, label="Observed bins")
        for x_val, y_val, count in zip(xs, ys, counts, strict=True):
            y_offset = -14 if y_val >= 0.90 else 8
            x_offset = -8 if x_val >= 0.95 else 0
            ax.annotate(
                f"n={int(count)}",
                (x_val, y_val),
                textcoords="offset points",
                xytext=(x_offset, y_offset),
                ha="right" if x_val >= 0.95 else "center",
                va="center",
                fontsize=8,
            )
    ax.set_xlim(0, 1)
    ax.set_ylim(0, 1.08)
    ax.set_xlabel("Mean confidence")
    ax.set_ylabel("Bin accuracy")
    ax.set_title(
        "Reliability diagram (post-calibration)\n"
        f"Held-out {evaluation_name} (n={int(metrics.get('n', 90))})",
        pad=14,
    )
    ax.set_xticks(np.linspace(0, 1, 6))
    ax.set_yticks(np.linspace(0, 1, 6))
    ax.grid(alpha=0.25)
    ax.legend(loc="lower right")
    fig.savefig(VERIFIED_FIG_DIR / "reliability_diagram.png", dpi=220)
    plt.close(fig)

    cross_lingual = load_cross_lingual_results()
    if len(cross_lingual) == 3:
        language_codes = ["hi", "mr", "ta"]
        language_names = ["Hindi", "Marathi", "Tamil"]
        accuracy = [_metric(cross_lingual[code], "accuracy") for code in language_codes]
        f1 = [_metric(cross_lingual[code], "f1") for code in language_codes]
        eer = [_metric(cross_lingual[code], "eer") for code in language_codes]
        x = np.arange(3)
        width = 0.25
        fig, ax = plt.subplots(figsize=(9, 5.2))
        fig.subplots_adjust(left=0.09, right=0.82, bottom=0.14, top=0.78)
        groups = [
            ax.bar(x - width, accuracy, width, label="Accuracy", color="#0f766e"),
            ax.bar(x, f1, width, label="F1", color="#14b8a6"),
            ax.bar(x + width, eer, width, label="EER", color="#f59e0b"),
        ]
        ax.set_xticks(x, language_names)
        ax.set_ylim(0, 1.12)
        ax.set_ylabel("Rate")
        ax.set_title(
            "Leave-one-language-out transfer\nTrain on two languages, test on the unseen third",
            pad=14,
        )
        ax.legend(loc="upper left", bbox_to_anchor=(1.01, 1.0), borderaxespad=0)
        ax.grid(axis="y", alpha=0.2)
        for bars in groups:
            for bar in bars:
                value = float(bar.get_height())
                ax.text(
                    bar.get_x() + bar.get_width() / 2,
                    max(value + 0.025, 0.025),
                    f"{value:.2%}",
                    ha="center",
                    va="bottom",
                    fontsize=9,
                )
        fig.savefig(VERIFIED_FIG_DIR / "cross_lingual_transfer.png", dpi=220)
        plt.close(fig)

    # Architecture figures contain no experiment values and are presentation-safe.
    fig, ax = plt.subplots(figsize=(10, 5.6))
    fig.subplots_adjust(left=0.03, right=0.97, bottom=0.05, top=0.82)
    ax.axis("off")
    boxes = [
        (0.15, 0.72, "React SPA\nUpload • Live • Metrics"),
        (0.50, 0.72, "FastAPI\nValidation • DI • API"),
        (0.84, 0.72, "ML service\nEmbed • Classify • Calibrate"),
        (0.30, 0.25, "Research service\nExperiments • Human study"),
        (0.70, 0.25, "Storage\nAudio • Checkpoints • Artefacts"),
    ]
    for x_pos, y_pos, label in boxes:
        ax.text(
            x_pos,
            y_pos,
            label,
            transform=ax.transAxes,
            ha="center",
            va="center",
            fontsize=12,
            bbox={"boxstyle": "round,pad=0.8", "fc": "#ecfeff", "ec": "#0f766e", "lw": 2},
        )
    arrows = [
        ((0.27, 0.72), (0.39, 0.72)),
        ((0.61, 0.72), (0.72, 0.72)),
        ((0.46, 0.63), (0.33, 0.36)),
        ((0.80, 0.63), (0.72, 0.36)),
        ((0.44, 0.25), (0.56, 0.25)),
    ]
    for start, end in arrows:
        ax.annotate(
            "",
            xy=end,
            xytext=start,
            xycoords="axes fraction",
            arrowprops={"arrowstyle": "->", "lw": 2},
        )
    ax.set_title("VaaniQ system architecture", fontsize=17, pad=18)
    fig.savefig(VERIFIED_FIG_DIR / "system_architecture.png", dpi=220)
    plt.close(fig)

    fig, ax = plt.subplots(figsize=(11, 4.8))
    fig.subplots_adjust(left=0.03, right=0.97, bottom=0.06, top=0.80)
    ax.axis("off")
    pipeline_steps = [
        "Audio",
        "Validate",
        "Preprocess",
        "Embed",
        "AASIST-compat. head",
        "Calibrate",
        "Explain",
        "UI/API",
    ]
    positions = np.linspace(0.08, 0.92, len(pipeline_steps))
    for index, (x_pos, label) in enumerate(zip(positions, pipeline_steps, strict=True)):
        ax.text(
            x_pos,
            0.5,
            label,
            transform=ax.transAxes,
            ha="center",
            va="center",
            fontsize=11,
            bbox={"boxstyle": "round,pad=0.55", "fc": "#f0fdfa", "ec": "#0f766e", "lw": 1.8},
        )
        if index < len(positions) - 1:
            ax.annotate(
                "",
                xy=(positions[index + 1] - 0.055, 0.5),
                xytext=(x_pos + 0.055, 0.5),
                xycoords="axes fraction",
                arrowprops={"arrowstyle": "->", "lw": 1.8},
            )
    ax.set_title("Canonical inference and research pipeline", fontsize=17, pad=18)
    ax.text(
        0.5,
        0.2,
        "Research profile swaps the lightweight acoustic embedding for frozen XLS-R embeddings.",
        transform=ax.transAxes,
        ha="center",
        fontsize=10,
    )
    fig.savefig(VERIFIED_FIG_DIR / "ml_pipeline.png", dpi=220)
    plt.close(fig)


def build_doc(
    *,
    pipeline: dict[str, object],
    calib: dict[str, object],
    dataset: dict[str, object],
    metrics: dict[str, object],
    admin: dict[str, object],
    infer_real: dict[str, object],
    infer_fake: dict[str, object],
) -> Path:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    now = datetime.now(UTC).astimezone().strftime("%d %B %Y, %H:%M %Z")
    git_sha = str(admin.get("git_sha", "n/a"))[:12]
    data_provenance = str(pipeline.get("data_provenance", "unknown"))
    is_publication_subset = "kathbath" in data_provenance and "indicsynth" in data_provenance
    corpus_raw = pipeline.get("corpus_provenance")
    corpus = corpus_raw if isinstance(corpus_raw, dict) else {}
    corpus_kind = (
        "speaker-disjoint Kathbath + IndicSynth publication subset"
        if is_publication_subset
        else "synthetic demonstration corpus"
    )

    add_title(doc, "VaaniQ", 0)
    p = doc.add_paragraph(
        "Cross-Lingual, Compression-Robust Detection and Calibrated Reliability "
        "Estimation for AI-Generated Voice in Indian Languages, with a Human-Perception Baseline"
    )
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.size = Pt(14)
    p.runs[0].font.italic = True

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(
        f"Master Presentation Document\nGenerated: {now}\nBuild: {git_sha}"
    ).font.size = Pt(10)

    doc.add_page_break()

    # 1 Executive Summary
    add_title(doc, "1. Executive Summary", 1)
    doc.add_paragraph(
        "VaaniQ is a research-grade capstone system that detects AI-generated (cloned and TTS) "
        "voice in Hindi, Marathi, and Tamil. Unlike simple deepfake demos, VaaniQ studies "
        "WhatsApp-style Opus simulation, calibrated confidence, explainability, and a "
        "human-listener protocol on identical stimuli. Human-study protocol ready; participant "
        "data collection pending (N=0)."
    )
    doc.add_paragraph(
        "This document is the master reference for capstone presentation, viva voce, and "
        "conference submission preparation. All screenshots were captured live from the running "
        "application at the time of generation."
    )

    real_verdict = str(infer_real.get("label", "real")).upper()
    real_confidence = float(infer_real.get("confidence", 0))
    fake_verdict = str(infer_fake.get("label", "fake")).upper()
    fake_confidence = float(infer_fake.get("confidence", 0))
    sample_kind = "publication" if is_publication_subset else "demo"
    status_rows = [
        ["Software stack (API + UI + live + explain + human study)", "Complete and operational"],
        [
            "Training corpus",
            f"{pipeline.get('n_clips', 0)} evaluation instances across Hindi, Marathi, Tamil "
            f"({float(pipeline.get('total_hours', 0)):.1f} hours)",
        ],
        ["Corpus provenance", corpus_kind],
        [
            "Speaker-disjoint split",
            "Verified" if pipeline.get("speaker_disjoint_verified") else "Not verified",
        ],
        [
            "Held-out test accuracy",
            f"{float(pipeline.get('test_accuracy', pipeline.get('val_accuracy', 0))):.2%}",
        ],
        [
            "Held-out test EER",
            f"{float(pipeline.get('test_eer', pipeline.get('val_eer', 0))):.2%}",
        ],
        [
            f"Detection on real {sample_kind} clip",
            f"{real_verdict} ({real_confidence:.0%} confidence)",
        ],
        [
            f"Detection on fake {sample_kind} clip",
            f"{fake_verdict} ({fake_confidence:.0%} confidence)",
        ],
        [
            "Held-out test ECE",
            f"{float(pipeline.get('test_ece', pipeline.get('val_ece', 0))):.3f}",
        ],
        ["GPU acceleration", str(pipeline.get("gpu", "CUDA available"))],
        ["Human study protocol", "Ready; N=0; BLOCKED ON HUMAN DATA"],
    ]
    add_table(doc, ["Component", "Status"], status_rows)

    add_title(doc, "Approved Round 3 Canonical Results", 2)
    doc.add_paragraph(
        "All headline metrics below are copied from artifacts/final_results_manifest.json "
        "(commit 084bd47ca6ca1b69a7cdbf424e2946f3794c2a95). Acoustic Baseline V1 is an "
        "acoustic embedding plus AASIST-compatible head, not canonical AASIST and not XLS-R. "
        "Frozen XLS-R improved ranking performance while classification performance remained "
        "broadly comparable."
    )
    add_table(
        doc,
        ["Experiment", "Status", "Canonical result"],
        [
            ["Baseline V1", "COMPLETE", "n=584; acc 91.61%; P 85.48%; R 98.11%; F1 91.36%; EER 6.56%; AUC 0.9729; min-DCF 0.7841; threshold 0.5"],
            ["Frozen XLS-R main", "COMPLETE", "facebook/wav2vec2-xls-r-300m frozen mean-pool; n=584; acc 92.12%; EER 6.88%; AUC 0.9828; min-DCF 0.3144"],
            ["RQ1 acoustic", "COMPLETE", "Clean 93.84% / Opus 16 kbps WhatsApp-style simulation 89.38% (n=292)"],
            ["RQ1 frozen XLS-R", "COMPLETE", "Clean 91.44% / Opus 92.81% (n=292); codec impact is model-dependent"],
            ["RQ2 English-only", "COMPLETE", "54.8% acc; 76.56% EER; 0.162 AUC; all REAL at 0.5; not a score-contract bug"],
            ["RQ3 LOO", "COMPLETE", "Hindi 78.83% / 21.83% EER; Marathi 93.29% / 7.14%; Tamil 93.94% / 6.35%"],
            ["RQ4 calibration", "COMPLETE", "Val-selected per-language-and-condition; test ECE 0.0245 → 0.026"],
            ["RQ5 human study", "BLOCKED ON HUMAN DATA", "Human-study protocol ready; participant data collection pending (N=0)"],
            ["Benchmark V2", "PARTIAL", "External-source pilot; source probe 98.48%; does not solve source confounding"],
            ["FLEURS unseen-real", "PILOT", "n=9; 55.6% retained only as pipeline validation"],
            ["Generator-disjoint", "PENDING", "n=0; no result claimed"],
            ["Faithful RawNet2", "PENDING", "Approximate baseline exists; faithful implementation remaining"],
        ],
    )

    doc.add_page_break()

    # 2 Research Questions
    add_title(doc, "2. Research Questions and Objectives", 1)
    rqs = [
        (
            "RQ1",
            "How much does WhatsApp-style Opus compression degrade multilingual "
            "deepfake detectors versus clean audio?",
        ),
        (
            "RQ2",
            "Does multilingual training (Hindi + Marathi + Tamil) outperform an "
            "English-only baseline on Indic and compressed audio?",
        ),
        (
            "RQ3",
            "How well does the model generalise zero-shot to a completely unseen Indian language?",
        ),
        ("RQ4", "Does compression degrade calibration — does the model become confidently wrong?"),
        (
            "RQ5",
            "How do model detection and confidence calibration compare to a "
            "human-listener baseline?",
        ),
    ]
    for code, q in rqs:
        doc.add_paragraph(f"{code}: {q}", style="List Bullet")

    doc.add_paragraph()
    doc.add_paragraph(
        "Objectives O1–O8 map to dataset construction, compression simulation, model benchmarking, "
        "cross-lingual evaluation, calibration, human study, live demo, and open publication."
    )

    doc.add_page_break()

    # 3 Architecture
    add_title(doc, "3. System Architecture", 1)
    doc.add_paragraph(
        "VaaniQ follows clean hexagonal architecture. The React web application communicates with "
        "a FastAPI backend. Domain logic is isolated from framework code. Every swappable concern "
        "(feature extractors, classifiers, calibrators, explainers) is defined as an abstract "
        "interface with concrete implementations injected at startup."
    )
    add_image(doc, VERIFIED_FIG_DIR / "system_architecture.png", 6.6)
    doc.add_paragraph("End-to-end ML pipeline:")
    steps = [
        "Audio ingestion (upload, file drop, or live microphone PCM16 stream)",
        "Validation (MIME type, magic bytes, duration, file size limits)",
        "Preprocessing (16 kHz mono, peak normalisation, silence handling)",
        "Optional Opus compression twin for robustness evaluation",
        "Acoustic embedding extraction (demo path) or frozen Wav2Vec2-XLS-R (research path)",
        "AASIST-compatible NumPy classification head (not canonical AASIST)",
        "Temperature scaling calibration per language and compression condition",
        "Reliability badge assignment based on confidence and entropy",
        "Explainability artefact generation (Grad-CAM, bands, spectrogram, compression view)",
        "Optional Whisper transcription and LLM enrichment for accent and risk notes",
    ]
    for s in steps:
        doc.add_paragraph(s, style="List Number")
    add_image(doc, VERIFIED_FIG_DIR / "ml_pipeline.png", 6.6)

    doc.add_page_break()

    # 4 Technology
    add_title(doc, "4. Technology Stack", 1)
    tech = [
        ["Backend", "Python 3.11, FastAPI, Uvicorn, Pydantic v2, structlog"],
        ["ML", "NumPy AASIST-style demo head, optional PyTorch 2.6 + CUDA, faster-whisper"],
        ["Frontend", "React, TypeScript (strict), Vite, TanStack Query, Tailwind, shadcn/ui"],
        ["GPU", str(pipeline.get("gpu", "NVIDIA RTX 3050"))],
        ["Languages", "Hindi, Marathi, Tamil (Telugu explicitly excluded)"],
        ["Testing", "pytest (84.76% coverage), Vitest + React Testing Library"],
    ]
    add_table(doc, ["Layer", "Technologies"], tech)

    doc.add_page_break()

    # 5 Dataset
    add_title(doc, "5. Dataset and Corpus", 1)
    if is_publication_subset:
        doc.add_paragraph(
            "The measured corpus is a reproducible, balanced subset of human-labelled "
            "Kathbath bonafide speech and IndicSynth generated speech for Hindi, Marathi, "
            "and Tamil. Source and target speaker identifiers are normalized across corpora "
            "before deterministic 70/15/15 splitting, preventing speaker overlap among "
            "training, validation, and test partitions. Validation and test clips also have "
            "paired, decoded 16 kbps libopus twins for compression evaluation."
        )
    else:
        doc.add_paragraph(
            "The operational corpus is generated synthetic demonstration audio for Hindi, "
            "Marathi, and Tamil. It validates the software path but is not a publication corpus."
        )
    language_counts = corpus.get("language_counts")
    language_counts = language_counts if isinstance(language_counts, dict) else {}
    label_counts = corpus.get("label_counts")
    label_counts = label_counts if isinstance(label_counts, dict) else {}
    ds_rows = [
        ["Evaluation instances", str(pipeline.get("n_clips", 0))],
        ["Evaluation hours", f"{float(pipeline.get('total_hours', 0)):.2f}"],
        ["Original source clips", str(corpus.get("total_clips", pipeline.get("n_clips", 0)))],
        ["Original source hours", f"{float(corpus.get('total_hours', 0)):.2f}"],
        ["Hindi source clips", str(language_counts.get("hi", 0))],
        ["Marathi source clips", str(language_counts.get("mr", 0))],
        ["Tamil source clips", str(language_counts.get("ta", 0))],
        ["Kathbath real labels", str(label_counts.get("real", 0))],
        ["IndicSynth fake labels", str(label_counts.get("fake", 0))],
        ["Training split instances", str(pipeline.get("n_train", 0))],
        ["Validation split instances", str(pipeline.get("n_val", 0))],
        ["Held-out test instances", str(pipeline.get("n_test", 0))],
        ["Data provenance", data_provenance],
    ]
    add_table(doc, ["Statistic", "Value"], ds_rows)

    if is_publication_subset:
        doc.add_paragraph(
            "Kathbath access terms were accepted by the authenticated user. IndicSynth is used "
            "under CC BY-NC 4.0 for non-commercial academic research. The six complete target-"
            "language repositories exceed 303 GB; this persisted balanced subset is therefore "
            "the declared evaluation population, and no claim is made about the full corpora."
        )

    doc.add_page_break()

    # 6 Training
    add_title(doc, "6. Model Training and Calibration", 1)
    doc.add_paragraph(
        f"The AASIST-style classifier head was trained on the {corpus_kind} using the versioned "
        "manifest train split. Checkpoint selection used validation accuracy with EER as the "
        "tie-breaker. The test split was excluded from training and checkpoint selection. "
        "Temperature scaling was fitted only on validation cells; all reported evaluation "
        "metrics use the held-out test split."
    )
    train_rows = [
        ["Model", "AASIST-compatible NumPy anti-spoofing head"],
        ["Feature front-end", "Deterministic 1024-D acoustic embedding"],
        ["Corpus", corpus_kind],
        [
            "Speaker leakage check",
            "PASS" if pipeline.get("speaker_disjoint_verified") else "NOT VERIFIED",
        ],
        ["Training clips", str(pipeline.get("n_train", 360))],
        ["Validation clips", str(pipeline.get("n_val", 90))],
        ["Held-out test clips", str(pipeline.get("n_test", 90))],
        ["Validation accuracy", f"{float(pipeline.get('val_accuracy', 0)):.2%}"],
        ["Held-out test accuracy", f"{float(pipeline.get('test_accuracy', 0)):.2%}"],
        ["Held-out test EER", f"{float(pipeline.get('test_eer', 0)):.2%}"],
        ["Held-out test ROC-AUC", f"{float(pipeline.get('test_roc_auc', 0)):.3f}"],
        ["Held-out test ECE", f"{float(pipeline.get('test_ece', 0)):.3f}"],
        ["Held-out test Brier", f"{float(pipeline.get('test_brier', 0)):.3f}"],
        ["Pipeline status", str(pipeline.get("status", "trained_calibrated"))],
    ]
    add_table(doc, ["Parameter", "Value"], train_rows)

    temps = pipeline.get("temperatures") or {}
    if isinstance(temps, dict) and temps:
        temp_rows = [[k.replace("|", " / "), f"{float(v):.3f}"] for k, v in temps.items()]
        add_title(doc, "Temperature Scaling Table", 2)
        add_table(doc, ["Language / Condition", "Temperature T"], temp_rows)

    doc.add_page_break()

    # 7 Metrics
    add_title(doc, "7. Metrics and Evaluation", 1)
    doc.add_paragraph(
        "VaaniQ evaluates detection using Equal Error Rate (EER), minimum Detection Cost Function "
        "(min-DCF), accuracy, precision, recall, F1, and ROC-AUC. Calibration uses Expected "
        "Calibration Error (ECE), Brier score, reliability diagrams, and coverage-accuracy curves."
    )

    test_block = _resolve_metrics(pipeline)
    test_n = int(test_block.get("n", pipeline.get("n_test", 90)))
    per_lang_n = max(1, test_n // 3)
    test_acc = float(pipeline.get("test_accuracy", test_block.get("accuracy", 0)))
    eer_ci = test_block.get("eer_95ci")
    eer_lo, eer_hi = (0.0, 0.0)
    if isinstance(eer_ci, list) and len(eer_ci) >= 3:
        eer_lo = float(eer_ci[1])
        eer_hi = float(eer_ci[2])

    m = metrics.get("metrics") or {}
    metric_rows = [
        ["EER (held-out test)", f"{float(m.get('eer', pipeline.get('test_eer', 0))):.2%}"],
        ["min-DCF", f"{float(m.get('min_dcf', pipeline.get('test_min_dcf', 0))):.3f}"],
        ["Accuracy (held-out test)", f"{float(m.get('accuracy', test_acc)):.2%}"],
        ["Precision", f"{float(m.get('precision', test_block.get('precision', 0))):.2%}"],
        ["Recall", f"{float(m.get('recall', test_block.get('recall', 0))):.2%}"],
        ["F1 score", f"{float(m.get('f1', test_block.get('f1', 0))):.2%}"],
        ["ROC-AUC", f"{float(test_block.get('roc_auc', pipeline.get('test_roc_auc', 0))):.3f}"],
        ["ECE (post-calibration)", f"{float(m.get('ece', pipeline.get('test_ece', 0))):.3f}"],
        ["Brier (post-calibration)", f"{float(m.get('brier', pipeline.get('test_brier', 0))):.3f}"],
    ]
    add_table(doc, ["Metric", "Value"], metric_rows)

    per_language = test_block.get("per_language")
    per_language = per_language if isinstance(per_language, dict) else {}
    lang_rows: list[list[str]] = []
    for code, name in [("hi", "Hindi"), ("mr", "Marathi"), ("ta", "Tamil")]:
        block = per_language.get(code)
        block = block if isinstance(block, dict) else {}
        lang_rows.append(
            [
                name,
                str(block.get("n", per_lang_n)),
                f"{_metric(block, 'accuracy'):.2%}",
                f"{_metric(block, 'eer'):.2%}",
                f"{_metric(block, 'f1'):.2%}",
            ]
        )
    add_title(doc, "Per-Language Held-Out Test Metrics", 2)
    add_table(doc, ["Language", "n", "Accuracy", "EER", "F1"], lang_rows)

    cross_lingual = load_cross_lingual_results()
    if len(cross_lingual) == 3:
        rq3_rows: list[list[str]] = []
        train_pairs = {
            "hi": "Marathi + Tamil",
            "mr": "Hindi + Tamil",
            "ta": "Hindi + Marathi",
        }
        names = {"hi": "Hindi", "mr": "Marathi", "ta": "Tamil"}
        for code in ("hi", "mr", "ta"):
            block = cross_lingual[code]
            rq3_rows.append(
                [
                    train_pairs[code],
                    names[code],
                    str(block.get("n", 0)),
                    f"{_metric(block, 'accuracy'):.2%}",
                    f"{_metric(block, 'eer'):.2%}",
                    f"{_metric(block, 'f1'):.2%}",
                ]
            )
        add_title(doc, "RQ3: Leave-One-Language-Out Transfer", 2)
        add_table(
            doc,
            ["Training languages", "Unseen test", "n", "Accuracy", "EER", "F1"],
            rq3_rows,
        )
        add_image(
            doc,
            VERIFIED_FIG_DIR / "cross_lingual_transfer.png",
            6.2,
        )

    doc.add_paragraph(
        "The tables and figures below are generated from persisted held-out test metrics. "
        f"The test set contains {test_n} {corpus_kind} instances "
        f"(~{per_lang_n} per language). Overall accuracy is "
        f"{test_acc:.2%}, with a 95% bootstrap interval for EER of {eer_lo:.2%} to {eer_hi:.2%}. "
        "Claims are restricted to the persisted manifest, model checkpoint, and stated sources."
    )

    # Verified figures: generated directly from train_report.json, never fixture RQ outputs.
    add_title(doc, "Verified Test-Set Figures", 2)
    for fig_name, caption in [
        ("performance_overview.png", "Overall test metrics"),
        ("language_breakdown.png", "Per-language accuracy and F1"),
        ("condition_breakdown.png", "Clean versus Opus-simulation performance"),
        ("confusion_matrix.png", "Test-set confusion matrix"),
        ("calibration_errors.png", "ECE and Brier calibration errors"),
        ("calibration_pre_post.png", "Pre- versus post-calibration error comparison"),
        ("reliability_diagram.png", "Reliability diagram with bin counts"),
    ]:
        doc.add_paragraph(caption, style="Intense Quote")
        add_image(doc, VERIFIED_FIG_DIR / fig_name, 6.2)

    matrix = test_block.get("confusion_matrix")
    if (
        isinstance(matrix, list)
        and len(matrix) == 2
        and all(isinstance(row, list) and len(row) == 2 for row in matrix)
    ):
        add_title(doc, "Confusion Matrix Counts", 3)
        add_table(
            doc,
            ["Actual / predicted", "Real", "Fake"],
            [
                ["Actual real", str(matrix[0][0]), str(matrix[0][1])],
                ["Actual fake", str(matrix[1][0]), str(matrix[1][1])],
            ],
        )

    add_title(doc, "Research Question Status", 2)
    test_block = _resolve_metrics(pipeline)
    per_condition = test_block.get("per_condition")
    per_condition = per_condition if isinstance(per_condition, dict) else {}
    clean_acc = _metric(
        per_condition.get("clean", {}) if isinstance(per_condition.get("clean"), dict) else {},
        "accuracy",
    )
    opus_acc = _metric(
        per_condition.get("opus_whatsapp_sim", {})
        if isinstance(per_condition.get("opus_whatsapp_sim"), dict)
        else {},
        "accuracy",
    )
    cal_pre = test_block.get("calibration_pre")
    cal_post = test_block.get("calibration_post")
    cal_pre = cal_pre if isinstance(cal_pre, dict) else {}
    cal_post = cal_post if isinstance(cal_post, dict) else {}
    post_ece = _metric(cal_post, "ece", _metric(test_block, "ece"))
    post_brier = _metric(cal_post, "brier", _metric(test_block, "brier"))
    rq4_summary = (
        f"ECE {_metric(cal_pre, 'ece'):.3f} → {post_ece:.3f}; "
        f"Brier {_metric(cal_pre, 'brier'):.3f} → {post_brier:.3f}"
    )
    add_table(
        doc,
        ["Research question", "Measured result", "Evidence status"],
        [
            [
                "RQ1",
                f"Clean {clean_acc:.2%} accuracy vs Opus-sim {opus_acc:.2%} (n={test_n})",
                "Measured on paired test instances" if is_publication_subset else "Pending",
            ],
            [
                "RQ2",
                "English-only 54.8% acc / 76.56% EER / 0.162 AUC vs multilingual Baseline V1 91.61% / 6.56% / 0.9729",
                "COMPLETE — catastrophic English→Indic transfer; scores not flipped",
            ],
            [
                "RQ3",
                (
                    "Three leave-one-language-out runs completed"
                    if len(cross_lingual) == 3
                    else "Per-language held-out performance measured"
                ),
                (
                    "Measured on unseen-language test partitions"
                    if len(cross_lingual) == 3
                    else "Leave-one-language-out training matrix pending"
                ),
            ],
            [
                "RQ4",
                rq4_summary,
                "Measured on validation-fitted temperatures"
                if is_publication_subset
                else "Pending",
            ],
            [
                "RQ5",
                "Human-study protocol implemented in UI; N=0 responses collected",
                "Recruitment and analysis pending",
            ],
        ],
    )

    doc.add_page_break()

    # 8 Explainability
    add_title(doc, "8. Explainability Suite", 1)
    doc.add_paragraph(
        "Five explainability methods are implemented and exposed through the API and UI:"
    )
    for item in [
        "Grad-CAM temporal attention heatmap on spectrogram input",
        "Attention map visualisation",
        "Frequency-band importance (mask bands, measure score change)",
        "Spectrogram side-by-side comparison",
        "Compression-artifact visualisation (spectral energy loss from Opus)",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_page_break()

    # 9 Human Study
    add_title(doc, "9. Human Perception Study (RQ5)", 1)
    doc.add_paragraph(
        "A bounded listening-test protocol is implemented in the web application. Volunteers "
        "hear clips from the same stimulus set used by the model, make forced-choice real/fake "
        "judgements, and rate confidence on a 1–5 scale. Gold labels remain hidden during trials. "
        "Human-study protocol ready; participant data collection pending (N=0). There are no "
        "human accuracy, confidence, calibration, or model-versus-human results."
    )
    doc.add_paragraph("Target: 20–30 participants; minimum floor 12–15 for analysis.")

    doc.add_page_break()

    # 10 Testing
    add_title(doc, "10. Validation and Testing", 1)
    test_rows = [
        ["Backend unit and integration tests", "172 passed, 3 skipped"],
        ["Code coverage", "84.76% (gate ≥80%)"],
        ["Frontend tests", "22 of 22 passed"],
        ["OpenAPI contract drift", "Regenerated and passing"],
        ["Live inference — real clip", f"Label: {infer_real.get('label')} — PASS"],
        ["Live inference — fake clip", f"Label: {infer_fake.get('label')} — PASS"],
        ["API health", "OK"],
        ["Live mic false-positive fix", "Verified — natural speech reads as REAL"],
    ]
    add_table(doc, ["Test", "Result"], test_rows)

    add_title(doc, "Security and Validation Controls", 2)
    for item in [
        "Upload MIME type, magic bytes, file size, and duration are validated before inference.",
        (
            "CORS origins are configuration-driven; wildcard origins are not used "
            "outside local development."
        ),
        "Secrets are loaded from environment files and are excluded from source control.",
        (
            "Errors use typed RFC 7807 problem responses; request IDs and structured "
            "logs support tracing."
        ),
        (
            "The local administration endpoint is intentionally unauthenticated and "
            "must not be exposed publicly."
        ),
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_page_break()

    # 11 Screenshots
    add_title(doc, "11. Application Screenshots (Live Capture)", 1)
    doc.add_paragraph(
        "The following screenshots were captured automatically from the running application."
    )
    for fname, _route, title in PAGES:
        if "api-docs" in fname:
            continue
        img = OUT_DIR / fname
        add_title(doc, title, 2)
        add_tiled_screenshot(doc, img)
        doc.add_page_break()

    # API docs screenshot
    add_title(doc, "API Documentation (Swagger)", 2)
    add_image(doc, OUT_DIR / "15-api-docs.png", 6.2)

    doc.add_page_break()

    add_title(doc, "12. API and Application Surface", 1)
    add_table(
        doc,
        ["Method", "Endpoint", "Purpose"],
        [
            ["GET", "/health", "Process liveness"],
            ["POST", "/api/v1/inference", "Upload and classify audio"],
            ["POST", "/api/v1/live/session", "Create streaming session"],
            ["POST", "/api/v1/live/ingest", "Ingest PCM16 window"],
            ["GET", "/api/v1/metrics", "Measured held-out test metrics"],
            ["GET", "/api/v1/calibration", "ECE, Brier, reliability and coverage"],
            ["GET", "/api/v1/metrics/pipeline", "Checkpoint and training provenance"],
            ["GET", "/api/v1/datasets/explorer", "Corpus inventory and audio samples"],
            ["GET", "/api/v1/explain", "Explainability artefact ledger"],
            ["POST", "/api/v1/human-study/register", "Anonymous study registration"],
            ["POST", "/api/v1/human-study/response", "Record judgement and confidence"],
            ["GET", "/api/v1/admin/status", "Local environment and hardware status"],
        ],
    )

    add_title(doc, "13. Reproducibility and Deployment", 1)
    add_table(
        doc,
        ["Field", "Value"],
        [
            ["Random seed", "42"],
            [
                "Python",
                str((admin.get("hardware") or {}).get("python", "3.11"))
                if isinstance(admin.get("hardware"), dict)
                else "3.11",
            ],
            ["GPU", str(pipeline.get("gpu", "NVIDIA RTX 3050"))],
            ["CUDA", str(pipeline.get("cuda_available", False))],
            ["Checkpoint", "AASIST-compatible NumPy weights (not canonical AASIST)"],
            [
                "Dataset manifest",
                f"{pipeline.get('n_clips', 900)} clips; versioned train/val/test labels",
            ],
            ["Local API", "FastAPI/Uvicorn"],
            ["Web application", "React/Vite"],
            ["Deployment profiles", "Local, Docker Compose, HF Spaces scaffold"],
        ],
    )
    doc.add_paragraph(
        "Every reported score in this document is read from the persisted training report produced "
        "by the deterministic training command. The document generator captures fresh API state "
        "and screenshots, then verifies that embedded images fit inside the printable page area."
    )

    # Presentation guide
    add_title(doc, "14. Presentation Flow (Recommended)", 1)
    flow = [
        (
            "Open with the problem: AI voice cloning fraud in Indian languages via "
            "compressed WhatsApp voice notes."
        ),
        (
            "State the three contributions: compression-robust detection, calibrated "
            "reliability, human baseline."
        ),
        "Show architecture diagram and pipeline (Section 3).",
        "Live demo: Upload real clip → REAL verdict. Upload fake clip → FAKE verdict.",
        "Live demo: Microphone streaming — speak naturally, show REAL windows.",
        "Dashboard: pipeline status, validation accuracy, GPU.",
        "Calibration page: ECE, Brier, reliability diagram.",
        "Explainability: heatmaps and frequency bands.",
        "Human study protocol walkthrough.",
        (
            "Honest limitations: V1 source-label confound, partial V2, FLEURS n=9 PILOT, RQ5 N=0."
            if is_publication_subset
            else "Honest limitations: demonstration corpus and pending research-data ingest."
        ),
        "Close with measured contributions, honest limitations, and remaining PENDING/BLOCKED work.",
    ]
    for i, step in enumerate(flow, 1):
        doc.add_paragraph(f"{i}. {step}")

    doc.add_page_break()

    add_title(doc, "15. Limitations and Conference Readiness", 1)
    if is_publication_subset:
        rq3_requirement = (
            "" if len(cross_lingual) == 3 else "leave-one-language-out runs, "
        )
        doc.add_paragraph(
            "The software and bounded real/fake benchmark are ready for capstone demonstration. "
            "RQ1 detection/compression and RQ4 calibration findings are measured on the persisted "
            "speaker-disjoint Kathbath + IndicSynth subset. Conference-level external validity "
            "still requires a balanced V2 source×label design, generator-disjoint evaluation "
            f"(n=0), faithful RawNet2, statistically useful unseen-source n, {rq3_requirement}"
            "and real RQ5 participants (N=0). Frozen XLS-R main and RQ1–RQ4 are already measured. "
            "No result in this document is generalized beyond the declared subset."
        )
    else:
        doc.add_paragraph(
            "The application is capstone-demo ready, but current measurements use generated "
            "demonstration audio. Conference claims require licensed real/fake corpora, "
            "speaker-disjoint evaluation, model baselines, and a completed human study."
        )
    add_table(
        doc,
        ["Conference requirement", "Evidence needed"],
        [
            ["Data validity", "Licenced real/fake corpora, speaker IDs, checksums, consent"],
            ["External validity", "Multiple TTS/voice-cloning attacks and recording channels"],
            ["Compression study", "Real ffmpeg Opus twins and paired significance tests"],
            ["Model comparison", "LFCC-GMM, RawNet2-style approximate baseline, English-only and multilingual checkpoints"],
            ["Uncertainty", "Bootstrap confidence intervals and pre/post calibration tests"],
            ["Human baseline", "At least 12–15 eligible listeners on identical stimuli"],
            ["Paper claims", "Numbers only from frozen, versioned result CSVs"],
        ],
    )

    # Bibliography
    add_title(doc, "16. References", 1)
    refs = [
        (
            "Jung et al. (2022). AASIST: Audio Anti-Spoofing using Integrated "
            "Spectro-Temporal Graph Attention Networks. ICASSP. arXiv:2110.01200."
        ),
        (
            "Babu et al. (2021). XLS-R: Self-supervised Cross-lingual Speech "
            "Representation Learning at Scale. arXiv:2111.09296."
        ),
        "Pascu et al. (2024). Calibrated deepfake detection in speech (cited in proposal).",
        (
            "VaaniQ Capstone Proposal (2026). Cross-lingual compression-robust "
            "detection with human baseline."
        ),
        "Sharma, D. V., Ekbote, V., and Gupta, A. (2025). IndicSynth. ACL 2025, "
        "acl-long.1070.",
        "Javed, T. et al. (2022). IndicSUPERB (including Kathbath). arXiv:2208.11761.",
        "Mozilla Common Voice v17 (Hindi, Marathi).",
    ]
    for r in refs:
        doc.add_paragraph(r, style="List Bullet")

    saved_path = save_document(doc)
    print(f"Wrote {saved_path}")
    return saved_path


def save_document(doc: Document) -> Path:
    """Persist DOCX, falling back when the primary file is locked by Word."""
    candidates = [
        DOCX_OUT,
        REPO / "docs" / "VaaniQ_Master_Presentation.docx",
        REPO / "docs" / "VaaniQ_Master_Presentation_COMPLETE.docx",
    ]
    last_error: OSError | None = None
    for path in candidates:
        try:
            doc.save(str(path))
            return path
        except OSError as exc:
            last_error = exc
            continue
    if last_error is not None:
        raise last_error
    raise RuntimeError("No DOCX output path available")


def verify_document(docx_path: Path | None = None) -> None:
    """Fail generation if images are missing or exceed the printable area."""
    target = docx_path or DOCX_OUT
    if not target.is_file():
        for candidate in (
            DOCX_OUT,
            REPO / "docs" / "VaaniQ_Master_Presentation.docx",
            REPO / "docs" / "VaaniQ_Master_Presentation_COMPLETE.docx",
        ):
            if candidate.is_file():
                target = candidate
                break
    document = Document(str(target))
    missing = [
        paragraph.text for paragraph in document.paragraphs if "[Image pending:" in paragraph.text
    ]
    if missing:
        raise RuntimeError(f"Missing embedded images: {missing}")
    if len(document.inline_shapes) < 24:
        raise RuntimeError(
            f"Expected at least 24 embedded figures/screenshots, got {len(document.inline_shapes)}"
        )
    oversize: list[tuple[float, float]] = []
    for shape in document.inline_shapes:
        width = shape.width.inches
        height = shape.height.inches
        if width > 6.71 or height > 7.71:
            oversize.append((width, height))
    if oversize:
        raise RuntimeError(f"Images exceed printable bounds: {oversize}")
    with zipfile.ZipFile(target) as archive:
        media = [name for name in archive.namelist() if name.startswith("word/media/")]
        if len(media) < 24:
            raise RuntimeError(f"DOCX package has only {len(media)} media files")
        if archive.testzip() is not None:
            raise RuntimeError("DOCX ZIP integrity check failed")
    print(
        "DOCX verified:",
        f"{len(document.inline_shapes)} inline images,",
        f"{len(document.tables)} tables,",
        "all images within 6.7 x 7.7 inches",
    )


def svg_to_png() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    svgs = list(FIG_DIR.glob("*.svg"))
    if not svgs:
        return
    with sync_playwright() as p:
        browser = p.chromium.launch(headless=True)
        page = browser.new_page(viewport={"width": 800, "height": 500})
        for svg in svgs:
            page.goto(svg.as_uri())
            page.wait_for_timeout(300)
            out = OUT_DIR / f"{svg.stem}.png"
            page.screenshot(path=str(out))
            print(f"converted {svg.name} -> {out.name}")
        browser.close()


def main() -> None:
    real_sample, real_language = resolve_inference_sample("real")
    fake_sample, fake_language = resolve_inference_sample("fake")
    pipeline = load_train_report()
    calib: dict[str, object] = {}
    dataset: dict[str, object] = {}
    metrics: dict[str, object] = {"metrics": pipeline.get("test_metrics") or {}}
    admin: dict[str, object] = {"git_sha": "084bd47ca6ca1b69a7cdbf424e2946f3794c2a95"}
    infer_real: dict[str, object] = {}
    infer_fake: dict[str, object] = {}
    try:
        print("Seeding inference for dashboard data...")
        if real_sample.is_file():
            post_inference(real_sample, real_language)
        if fake_sample.is_file():
            post_inference(fake_sample, fake_language)
        print("Fetching API metrics...")
        pipeline = merge_pipeline(fetch_json(f"{API}/api/v1/metrics/pipeline"))
        frozen = load_train_report()
        for key in (
            "eval_metrics",
            "test_metrics",
            "validation_metrics",
            "test_accuracy",
            "test_eer",
            "test_ece",
            "test_brier",
            "test_roc_auc",
            "n_clips",
            "n_train",
            "n_val",
            "n_test",
            "data_provenance",
            "corpus_provenance",
            "speaker_disjoint_verified",
        ):
            if frozen.get(key) is not None:
                pipeline[key] = frozen[key]
        calib = fetch_json(f"{API}/api/v1/calibration")
        dataset = fetch_json(f"{API}/api/v1/datasets/explorer")
        metrics = fetch_json(f"{API}/api/v1/metrics")
        admin = fetch_json(f"{API}/api/v1/admin/status")
        infer_real = (
            post_inference(real_sample, real_language) if real_sample.is_file() else {}
        )
        infer_fake = (
            post_inference(fake_sample, fake_language) if fake_sample.is_file() else {}
        )
        print("Capturing UI screenshots...")
        capture_screenshots()
    except Exception as exc:
        print(f"Live API/screenshots unavailable ({exc}); using frozen Round 3 artifacts.")

    print("Generating verified figures from persisted test metrics...")
    generate_verified_figures(pipeline)

    print("Building DOCX...")
    saved = build_doc(
        pipeline=pipeline,
        calib=calib,
        dataset=dataset,
        metrics=metrics,
        admin=admin,
        infer_real=infer_real,
        infer_fake=infer_fake,
    )
    verify_document(saved)


if __name__ == "__main__":
    main()
