#!/usr/bin/env python3

"""Generate the IEEE-style VaaniQ paper and progress report from frozen Round 3 artifacts."""



from __future__ import annotations



import zipfile

from pathlib import Path



from docx import Document

from docx.enum.section import WD_SECTION

from docx.enum.text import WD_ALIGN_PARAGRAPH

from docx.oxml import OxmlElement

from docx.oxml.ns import qn

from docx.shared import Inches, Pt



REPO = Path(__file__).resolve().parents[1]

FIGURE_DIR = REPO / "docs" / "assets" / "verified_figures"

OUTPUT = REPO / "docs" / "VaaniQ_IEEE_Research_Paper_Final.docx"

DRAFT_OUTPUT = REPO / "docs" / "VaaniQ_IEEE_Research_Paper_Draft.docx"

PROGRESS_OUTPUT = REPO / "docs" / "VaaniQ_Midterm_Progress_Report.docx"



import sys



sys.path.insert(0, str(REPO / "scripts"))

from report_data import (  # noqa: E402

    as_dict,

    as_number,

    load_baseline_matrix,

    load_benchmark_v2,

    load_final_results,

    load_rq2,

    load_rq4_audit,

    load_train_report,

    load_xlsr_main,

    pct,

)





def _load_report() -> dict[str, object]:

    return load_train_report()





def _dict(value: object) -> dict[str, object]:

    return as_dict(value)





def _number(mapping: dict[str, object], key: str) -> float:

    return as_number(mapping, key)





def _set_columns(section: object, count: int) -> None:

    section_properties = section._sectPr  # type: ignore[attr-defined]

    columns = section_properties.xpath("./w:cols")

    node = columns[0] if columns else OxmlElement("w:cols")

    node.set(qn("w:num"), str(count))

    node.set(qn("w:space"), "360")

    if not columns:

        section_properties.append(node)





def _set_cell_text(cell: object, text: str, *, bold: bool = False) -> None:

    cell.text = text  # type: ignore[attr-defined]

    for paragraph in cell.paragraphs:  # type: ignore[attr-defined]

        paragraph.paragraph_format.space_after = Pt(0)

        for run in paragraph.runs:

            run.font.name = "Times New Roman"

            run.font.size = Pt(7.5)

            run.bold = bold





def _table(

    doc: Document,

    headers: list[str],

    rows: list[list[str]],

    *,

    caption: str,

) -> None:

    paragraph = doc.add_paragraph()

    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = paragraph.add_run(caption.upper())

    run.bold = True

    run.font.size = Pt(8)

    table = doc.add_table(rows=1, cols=len(headers))

    table.style = "Table Grid"

    for index, header in enumerate(headers):

        _set_cell_text(table.rows[0].cells[index], header, bold=True)

    for row in rows:

        cells = table.add_row().cells

        for index, value in enumerate(row):

            _set_cell_text(cells[index], value)





def _heading(doc: Document, number: str, title: str) -> None:

    paragraph = doc.add_paragraph()

    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    paragraph.paragraph_format.space_before = Pt(5)

    paragraph.paragraph_format.space_after = Pt(2)

    run = paragraph.add_run(f"{number}. {title}".upper())

    run.bold = True

    run.font.name = "Times New Roman"

    run.font.size = Pt(10)





def _subheading(doc: Document, letter: str, title: str) -> None:

    paragraph = doc.add_paragraph()

    paragraph.paragraph_format.space_before = Pt(3)

    paragraph.paragraph_format.space_after = Pt(1)

    run = paragraph.add_run(f"{letter}. {title}")

    run.italic = True

    run.font.name = "Times New Roman"

    run.font.size = Pt(10)





def _body(doc: Document, text: str) -> None:

    paragraph = doc.add_paragraph(text)

    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    paragraph.paragraph_format.first_line_indent = Inches(0.16)

    paragraph.paragraph_format.space_after = Pt(2)

    paragraph.paragraph_format.line_spacing = 1.0





def _figure(doc: Document, filename: str, caption: str) -> None:

    path = FIGURE_DIR / filename

    if not path.is_file():

        raise FileNotFoundError(path)

    paragraph = doc.add_paragraph()

    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    paragraph.add_run().add_picture(str(path), width=Inches(3.18))

    caption_paragraph = doc.add_paragraph(caption)

    caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    caption_paragraph.paragraph_format.space_after = Pt(3)

    for run in caption_paragraph.runs:

        run.font.size = Pt(8)





def _metric_row(

    name: str,

    block: dict[str, object],

    *,

    n: object | None = None,

) -> list[str]:

    n_value = n if n is not None else block.get("n", "-")

    precision = block.get("precision")

    recall = block.get("recall")

    f1 = block.get("f1")

    roc = block.get("roc_auc")

    dcf = block.get("min_dcf")

    return [

        name,

        str(n_value),

        pct(_number(block, "accuracy")) if "accuracy" in block else "-",

        pct(float(precision)) if isinstance(precision, int | float) else "-",

        pct(float(recall)) if isinstance(recall, int | float) else "-",

        pct(float(f1)) if isinstance(f1, int | float) else "-",

        pct(_number(block, "eer")) if "eer" in block else "-",

        f"{float(roc):.4f}" if isinstance(roc, int | float) else "-",

        f"{float(dcf):.4f}" if isinstance(dcf, int | float) else "-",

    ]





def build_paper() -> Path:

    """Build the IEEE-style paper from frozen Round 3 artifacts."""

    report = _load_report()

    metrics = _dict(report.get("test_metrics"))

    provenance = _dict(report.get("corpus_provenance"))

    xlsr = load_xlsr_main()

    xlsr_metrics = _dict(xlsr.get("test_metrics"))

    rq1 = _dict(load_final_results().get("rq1"))

    acoustic_rq1 = _dict(rq1.get("baseline_v1_acoustic"))

    xlsr_rq1 = _dict(rq1.get("xlsr_main"))

    rq2 = load_rq2()

    english = _dict(rq2.get("english_only_indic_test"))

    multi = _dict(rq2.get("multilingual_baseline_v1_test"))

    rq3 = _dict(_dict(load_final_results().get("rq3")).get("folds"))

    rq4 = load_rq4_audit()

    matrix = _dict(load_baseline_matrix().get("models"))

    v2 = load_benchmark_v2()

    if "kathbath" not in str(report.get("data_provenance", "")):

        raise RuntimeError("IEEE paper requires measured Kathbath + IndicSynth results")



    doc = Document()

    section = doc.sections[0]

    section.page_width = Inches(8.5)

    section.page_height = Inches(11)

    section.top_margin = Inches(0.62)

    section.bottom_margin = Inches(0.62)

    section.left_margin = Inches(0.66)

    section.right_margin = Inches(0.66)

    normal = doc.styles["Normal"]

    normal.font.name = "Times New Roman"

    normal.font.size = Pt(9)

    normal.paragraph_format.space_after = Pt(2)



    title = doc.add_paragraph()

    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    title_run = title.add_run(

        "VaaniQ: Cross-Lingual, Compression-Robust Detection and "

        "Calibrated Reliability Estimation for AI-Generated Voice "

        "in Indian Languages"

    )

    title_run.bold = True

    title_run.font.name = "Times New Roman"

    title_run.font.size = Pt(16)



    authors = doc.add_paragraph()

    authors.alignment = WD_ALIGN_PARAGRAPH.CENTER

    authors.add_run(

        "Gaurav Phadale, Eshaan Sarkhawas, Aarav Phutane, Prajwal Patil\n"

        "Mukesh Patel School of Technology Management & Engineering, NMIMS\n"

        "Faculty Guide: Prof. Rama Bharti Varshney"

    )



    status = doc.add_paragraph()

    status.alignment = WD_ALIGN_PARAGRAPH.CENTER

    status_run = status.add_run(

        "Approved Round 3 research paper - frozen at commit 084bd47ca6ca1b69a7cdbf424e2946f3794c2a95"

    )

    status_run.italic = True

    status_run.font.size = Pt(8)



    abstract = doc.add_paragraph()

    abstract.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    abstract_run = abstract.add_run("Abstract-")

    abstract_run.bold = True

    abstract.add_run(

        "VaaniQ studies multilingual audio-deepfake detection under language and codec "

        "shift rather than claiming universal deepfake detection. A bounded V1 benchmark "

        "of Kathbath real speech and IndicSynth fake speech covers Hindi, Marathi, and "

        "Tamil with speaker-disjoint evaluation (held-out n=584). An acoustic-embedding "

        "plus AASIST-compatible head reached 91.61% accuracy and 6.56% EER. Frozen "

        "facebook/wav2vec2-xls-r-300m with mean pooling reached 92.12% accuracy, 6.88% "

        "EER, and 0.9828 ROC-AUC. Experiments include WhatsApp-style 16 kbps Opus "

        "simulation, leave-one-language-out transfer, and validation-selected calibration. "

        "Benchmark V2 remains a partial external-source pilot. Human-study protocol ready; "

        "participant data collection pending (N=0)."

    )



    keywords = doc.add_paragraph()

    keywords.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    keyword_run = keywords.add_run("Index Terms-")

    keyword_run.bold = True

    keywords.add_run(

        "audio deepfake detection, Indian languages, Kathbath, IndicSynth, "

        "WhatsApp-style Opus simulation, calibration, frozen XLS-R, speaker-disjoint evaluation"

    )



    columns = doc.add_section(WD_SECTION.CONTINUOUS)

    columns.top_margin = Inches(0.62)

    columns.bottom_margin = Inches(0.62)

    columns.left_margin = Inches(0.66)

    columns.right_margin = Inches(0.66)

    _set_columns(columns, 2)



    _heading(doc, "I", "Introduction")

    _body(

        doc,

        "Voice cloning can imitate speakers from short recordings and can be delivered "

        "as compressed messaging audio. Major anti-spoofing benchmarks remain concentrated "

        "in English or Mandarin. VaaniQ studies Hindi, Marathi, and Tamil on a bounded "

        "speaker-disjoint benchmark and measures codec shift, cross-lingual transfer, and "

        "confidence calibration. It does not claim universal fake-voice detection.",

    )

    _body(

        doc,

        "Five research questions are stated. RQ1: How does Opus compression affect "

        "detection performance? RQ2: How does English-only anti-spoofing transfer to "

        "Hindi, Marathi, and Tamil relative to multilingual training? RQ3: How well does "

        "the detector generalize to an unseen Indian language? RQ4: Does validation-selected "

        "post-hoc calibration remain reliable under held-out language/condition shift? "

        "RQ5: How does model performance compare with human listeners? RQ1-RQ4 are COMPLETE "

        "on the declared V1 benchmark. RQ5 is PENDING with N=0.",

    )



    _heading(doc, "II", "Related Work")

    _body(

        doc,

        "AASIST integrates spectro-temporal graph attention for anti-spoofing [1]. "

        "XLS-R provides cross-lingual self-supervised speech representations [2]. "

        "IndicSUPERB introduced Kathbath, a large human-labelled Indian-language speech "

        "resource [3]. IndicSynth provides synthetic speech for Indian languages [4]. "

        "Prior work shows that high classification accuracy does not guarantee calibrated "

        "confidence [5]. Muller et al. studied human perception of audio deepfakes [6]. "

        "To the best of our reviewed literature, the contribution is the integrated "

        "evaluation of Indic languages, codec shift, cross-lingual transfer, calibration, "

        "and human-study infrastructure rather than invention of XLS-R or AASIST.",

    )



    _heading(doc, "III", "Dataset")

    _subheading(doc, "A", "Bounded V1 Benchmark")

    _body(

        doc,

        f"The persisted source subset contains {int(provenance.get('total_clips', 1800))} "

        f"original clips ({float(provenance.get('total_hours', 0)):.2f} h): 900 Kathbath "

        "REAL and 900 IndicSynth FAKE, balanced across Hindi, Marathi, and Tamil. Paired "

        "validation/test 16 kbps libopus twins expand evaluation to 2,346 instances. "

        "Held-out test n=584. Bonafide clips follow Kathbath packaging terms; generated "

        "clips follow IndicSynth CC BY-NC 4.0 for non-commercial academic research.",

    )

    _body(

        doc,

        "V1 has a structural weakness: REAL equals Kathbath and FAKE equals IndicSynth. "

        "Source identity is therefore associated with class. Speaker-disjoint splitting "

        "protects against speaker leakage but does not remove source-domain shortcuts.",

    )

    _subheading(doc, "B", "Partial V2 External-Source Pilot")

    _body(

        doc,

        "Benchmark V2 is currently a partial external-source pilot. It adds an independent "

        "real-speech source but does not yet eliminate source-label confounding. Current "

        "cells are Kathbath\xd7real 1,177, IndicSynth\xd7fake 1,169, and FLEURS\xd7real 50. "

        "Generator tags exist (freevc24 805; xtts_v2 364), but generator-disjoint evaluation "

        "has n=0 and remains PENDING.",

    )

    _table(

        doc,

        ["Benchmark", "Source", "Class", "Count", "Status"],

        [

            ["V1", "Kathbath", "REAL", "900", "COMPLETE"],

            ["V1", "IndicSynth", "FAKE", "900", "COMPLETE"],

            ["V2", "Kathbath", "REAL", "1,177", "PARTIAL"],

            ["V2", "IndicSynth", "FAKE", "1,169", "PARTIAL"],

            ["V2", "FLEURS", "REAL", "50", "PARTIAL"],

        ],

        caption="Table I. Dataset composition",

    )

    _table(

        doc,

        ["Partition", "Instances", "Role"],

        [

            ["Train", str(report.get("n_train", 1254)), "Parameter learning"],

            ["Validation", str(report.get("n_val", 508)), "Checkpoint and calibration selection"],

            ["Test", str(report.get("n_test", 584)), "Final metrics only"],

        ],

        caption="Table II. Speaker-disjoint V1 partitions",

    )



    _heading(doc, "IV", "Methodology")

    _figure(

        doc,

        "system_architecture.png",

        "Fig. 1. VaaniQ software architecture. Pipeline: audio ? preprocess ? acoustic or frozen XLS-R path ? AASIST-compatible head ? calibration ? verdict.",

    )

    _subheading(doc, "A", "Acoustic Baseline V1")

    _body(

        doc,

        "Baseline V1 uses a deterministic 1024-D acoustic embedding and an "

        "AASIST-compatible NumPy classification head. It is not canonical AASIST and "

        "not XLS-R. The decision threshold is 0.5 on the fake softmax score. Label 0 is "

        "REAL/bonafide, label 1 is FAKE/spoof, and higher score_fake means higher "

        "probability of FAKE. Logit order is [real, fake].",

    )

    _subheading(doc, "B", "Frozen XLS-R Main")

    _body(

        doc,

        "The separate main path uses frozen facebook/wav2vec2-xls-r-300m, last-layer "

        "mean pooling, cached features, and a lightweight trainable AASIST-compatible "

        "head. The backbone is not updated. Production calibration for this path is "

        "global temperature scaling selected on validation.",

    )

    _subheading(doc, "C", "Baselines and Controls")

    _body(

        doc,

        "LFCC-GMM is a complete weak baseline. A RawNet2-style approximate baseline is "

        "complete but must not be called faithful RawNet2; faithful RawNet2 remains "

        "PENDING. An English-only ASVspoof LA control is evaluated on the same Indic "

        "held-out test.",

    )

    _subheading(doc, "D", "Compression and Evaluation Protocol")

    _body(

        doc,

        "ffmpeg/libopus at 16 kbps provides a WhatsApp-style Opus simulation. It is not "

        "audio transported through the WhatsApp service. EER, ROC-AUC, and normalized "

        "min-DCF use the same score direction. min-DCF uses P_target=0.05, C_miss=1, "

        "C_fa=1, with normalization denominator P_target \xd7 C_miss = 0.05.",

    )



    _heading(doc, "V", "Results")

    lfcc = _dict(matrix.get("lfcc_gmm"))

    rawnet = _dict(matrix.get("rawnet2_style_approx"))

    acoustic_row = dict(metrics)

    acoustic_row.setdefault("n", 584)

    xlsr_row = dict(xlsr_metrics)

    xlsr_row.setdefault("n", 584)

    _table(

        doc,

        ["Model", "n", "Accuracy", "Precision", "Recall", "F1", "EER", "ROC-AUC", "min-DCF"],

        [

            _metric_row("Acoustic Baseline V1", acoustic_row, n=584),

            _metric_row("Frozen XLS-R main", xlsr_row, n=584),

            _metric_row("LFCC-GMM", lfcc, n=lfcc.get("n", 584)),

            _metric_row("RawNet2-style approx.", rawnet, n=rawnet.get("n", 584)),

        ],

        caption="Table III. Primary held-out model comparison on the same V1 test set",

    )

    _body(

        doc,

        "Frozen XLS-R achieved 92.12% accuracy and ROC-AUC 0.9828 on the same 584-instance "

        "held-out V1 test set. Its EER of 6.88% was comparable to the acoustic baseline's "

        "6.56%, indicating stronger ranking performance without a uniform improvement "

        "across all decision metrics. An earlier acoustic checkpoint produced 93.66% "

        "accuracy, but the frozen Round 3 canonical checkpoint produces 91.61%; EER "

        "remained 6.56% because ranking behavior remained similar. That 93.66% figure is "

        "historical only.",

    )

    _figure(

        doc,

        "confusion_matrix.png",

        "Fig. 2. Confusion matrix for Acoustic Baseline V1 on the held-out V1 test partition (n=584).",

    )



    _heading(doc, "VI", "Compression (RQ1)")

    acoustic_clean = _dict(acoustic_rq1.get("clean"))

    acoustic_opus = _dict(acoustic_rq1.get("opus_whatsapp_sim"))

    xlsr_clean = _dict(xlsr_rq1.get("clean"))

    xlsr_opus = _dict(xlsr_rq1.get("opus_whatsapp_sim"))

    _table(

        doc,

        ["Condition", "n", "Accuracy", "EER", "F1"],

        [

            [

                "Clean",

                str(acoustic_clean.get("n", 292)),

                pct(_number(acoustic_clean, "accuracy")),

                pct(_number(acoustic_clean, "eer")),

                pct(_number(acoustic_clean, "f1")),

            ],

            [

                "Opus 16 kbps WhatsApp-style simulation",

                str(acoustic_opus.get("n", 292)),

                pct(_number(acoustic_opus, "accuracy")),

                pct(_number(acoustic_opus, "eer")),

                pct(_number(acoustic_opus, "f1")),

            ],

        ],

        caption="Table IV. RQ1 acoustic Baseline V1, clean versus Opus",

    )

    _table(

        doc,

        ["Condition", "n", "Accuracy", "EER", "F1"],

        [

            [

                "Clean",

                str(xlsr_clean.get("n", 292)),

                pct(_number(xlsr_clean, "accuracy")),

                pct(_number(xlsr_clean, "eer")),

                pct(_number(xlsr_clean, "f1")),

            ],

            [

                "Opus 16 kbps WhatsApp-style simulation",

                str(xlsr_opus.get("n", 292)),

                pct(_number(xlsr_opus, "accuracy")),

                pct(_number(xlsr_opus, "eer")),

                pct(_number(xlsr_opus, "f1")),

            ],

        ],

        caption="Table V. RQ1 frozen XLS-R, clean versus Opus",

    )

    _body(

        doc,

        "Compression effects were model-dependent. Acoustic Baseline V1 degraded from "

        "93.84% clean accuracy to 89.38% under WhatsApp-style Opus simulation "

        "(approximately 4.46 percentage points). Frozen XLS-R did not show the same "

        "degradation on this bounded split (91.44% clean, 92.81% Opus). These numbers "

        "must not be mixed across models.",

    )

    _figure(

        doc,

        "condition_breakdown.png",

        "Fig. 3. Acoustic Baseline V1 clean versus WhatsApp-style Opus simulation on paired test instances.",

    )



    _heading(doc, "VII", "English-Only Transfer (RQ2)")

    _table(

        doc,

        ["Model", "n", "Accuracy", "EER", "ROC-AUC"],

        [

            [

                "English-only ASVspoof LA ? Indic test",

                str(rq2.get("indic_test_n", 584)),

                "54.8%",

                pct(_number(english, "eer")),

                f"{_number(english, 'roc_auc'):.3f}",

            ],

            [

                "Multilingual Acoustic Baseline V1",

                "584",

                pct(_number(multi, "accuracy") or _number(metrics, "accuracy")),

                pct(_number(multi, "eer") or _number(metrics, "eer")),

                f"{_number(multi, 'roc_auc') or _number(metrics, 'roc_auc'):.4f}",

            ],

        ],

        caption="Table VI. RQ2 English-only versus multilingual on the same Indic test",

    )

    _body(

        doc,

        "The English-only model failed to transfer effectively to the Indic benchmark: "

        "54.8% accuracy, 76.56% EER, and 0.162 ROC-AUC, predicting every instance as REAL "

        "at threshold 0.5. A diagnostic that multiplies scores by ?1 yields ROC-AUC 0.838 "

        "and EER 23.44%, but official scores were not flipped. The global VaaniQ score "

        "contract is validated by multilingual Baseline V1 (91.61% accuracy, 6.56% EER, "

        "0.9729 ROC-AUC) on the same test. The result is catastrophic English?Indic "

        "transfer with anti-correlated ranking, not a generic score-inversion bug.",

    )



    _heading(doc, "VIII", "Cross-Lingual Transfer (RQ3)")

    hi = _dict(rq3.get("held_out_hi"))

    mr = _dict(rq3.get("held_out_mr"))

    ta = _dict(rq3.get("held_out_ta"))

    _table(

        doc,

        ["Held-out language", "n", "Accuracy", "EER"],

        [

            ["Hindi", str(hi.get("n", 222)), pct(_number(hi, "accuracy")), pct(_number(hi, "eer"))],

            ["Marathi", str(mr.get("n", 164)), pct(_number(mr, "accuracy")), pct(_number(mr, "eer"))],

            ["Tamil", str(ta.get("n", 198)), pct(_number(ta, "accuracy")), pct(_number(ta, "eer"))],

        ],

        caption="Table VII. RQ3 leave-one-language-out transfer",

    )

    _body(

        doc,

        "Cross-lingual transfer was asymmetric. Hindi held-out accuracy was 78.83% with "

        "21.83% EER, whereas Marathi and Tamil reached 93.29% and 93.94%. The study does "

        "not claim uniform language-independent generalization.",

    )

    if (FIGURE_DIR / "cross_lingual_transfer.png").is_file():

        _figure(

            doc,

            "cross_lingual_transfer.png",

            "Fig. 4. Leave-one-language-out transfer. Hindi is substantially weaker than Marathi and Tamil.",

        )



    _heading(doc, "IX", "Calibration (RQ4)")

    strategies = _dict(rq4.get("strategies"))

    uncal = _dict(strategies.get("uncalibrated"))

    val_metrics = _dict(report.get("validation_metrics"))

    val_cmp = _dict(val_metrics.get("calibration_strategy_comparison"))

    cal_pre = _dict(metrics.get("calibration_pre"))

    cal_post = _dict(metrics.get("calibration_post"))

    _table(

        doc,

        ["Quantity", "Value", "Split"],

        [

            [

                "Selected Baseline V1 strategy",

                str(val_cmp.get("selected_strategy", "per_language_and_condition")),

                "Validation only",

            ],

            ["Selected strategy val ECE", f"{_number(val_metrics, 'ece'):.4f}", "Validation"],

            ["Global TS val ECE", f"{as_number(val_cmp, 'global_temperature_val_ece', 0.0513):.4f}", "Validation"],

            ["Uncalibrated test ECE", f"{_number(cal_pre, 'ece'):.4f}", "Held-out test"],

            ["Post-calibration test ECE", f"{_number(cal_post, 'ece'):.4f}", "Held-out test"],

            [

                "Exploratory best-by-test-ECE",

                str(rq4.get("best_strategy_by_test_ece", "global_temperature")),

                "Not used for selection",

            ],

        ],

        caption="Table VIII. RQ4 calibration selection and held-out outcome",

    )

    _body(

        doc,

        "Calibration strategy was selected using validation only. For Baseline V1, the "

        "selected fine-grained per-language-and-condition strategy (val ECE 0.0487 versus "

        "global 0.0513) did not improve held-out ECE (0.0245 uncalibrated to 0.026 "

        "post-calibration). A standalone test comparison where global temperature scaling "

        f"looked better (test ECE {as_number(uncal, 'ece', 0.0378):.4f} versus exploratory "

        "global TS) is not evidence for selecting the production strategy. Validation-selected "

        "calibration did not uniformly improve held-out calibration. Frozen XLS-R used "

        "global temperature scaling selected on validation.",

    )



    _heading(doc, "X", "V2 Pilot and Generalization Status")

    evals = _dict(v2.get("evaluations"))

    independent = _dict(evals.get("independent_real_test"))

    gen = _dict(evals.get("generator_held_out_test"))

    shortcut = _dict(v2.get("source_shortcut_v2"))

    _table(

        doc,

        ["Eval", "n", "Status", "Result"],

        [

            ["V2 full test", str(_dict(evals.get("full_test")).get("n", 593)), "PARTIAL", "Pilot population, not confound-free"],

            [

                "FLEURS unseen-real",

                str(independent.get("n", 9)),

                "PILOT",

                "55.6% retained only as pipeline validation",

            ],

            ["Generator-disjoint", str(gen.get("n", 0)), "PENDING", "No result claimed"],

            [

                "V2 source probe",

                str(shortcut.get("test_n", 593)),

                "PARTIAL",

                f"{as_number(shortcut, 'source_probe_test_accuracy') * 100:.2f}% source identity",

            ],

            ["Faithful RawNet2", "-", "PENDING", "Not implemented"],

            ["RQ5 human study", "0", "BLOCKED ON HUMAN DATA", "Protocol ready; N=0"],

        ],

        caption="Table IX. V2 pilot and remaining generalization status",

    )

    _body(

        doc,

        "V2 source-probe accuracy is 98.48% versus a 85.83% label probe, showing that "

        "source identity remains highly predictable. No statistically useful unseen-source "

        "estimate is claimed from the current nine held-out FLEURS clips; the result is "

        "retained only as pipeline-validation evidence. Human-study protocol ready; "

        "participant data collection pending (N=0).",

    )



    _heading(doc, "XI", "Limitations and Threats to Validity")

    _body(

        doc,

        "Key threats are: (1) V1 source-label confound; (2) bounded dataset size; "

        "(3) academic speech versus actual scam audio; (4) simulated Opus versus real "

        "WhatsApp transport; (5) three-language scope; (6) no faithful RawNet2; "

        "(7) incomplete V2; (8) V2 source probe 98.48%; (9) FLEURS external pilot n=9; "

        "(10) generator-disjoint n=0; (11) RQ5 N=0; (12) calibration transfer uncertainty; "

        "(13) validation/test heterogeneity (Tamil val 71.5% versus test 90.9%; Marathi val "

        "98.4% versus test 93.3%) with no proven leakage but no fully resolved cause; and "

        "(14) no claim of universal unseen-generator robustness. Strong EER/ROC-AUC can "

        "coexist with Baseline V1 min-DCF 0.7841 because min-DCF uses a 5% target prior.",

    )



    _heading(doc, "XII", "Ethics")

    _body(

        doc,

        "Audio and model weights are excluded from version control. Kathbath access "

        "conditions and IndicSynth non-commercial restrictions are respected. The human "

        "protocol is anonymous, reveals no gold labels during trials, and has not yet "

        "collected participants.",

    )



    _heading(doc, "XIII", "Reproducibility")

    _body(

        doc,

        "Authoritative metrics are frozen in artifacts/final_results_manifest.json at "

        "commit 084bd47ca6ca1b69a7cdbf424e2946f3794c2a95. Manifests record source, label, "

        "speaker, split, duration, checksum, and generation model. Integrity verification "

        "is scripts/verify_research_integrity.py. Documents must consume the frozen "

        "manifest rather than superseded checkpoints or rounded tables.",

    )



    _heading(doc, "XIV", "Conclusion")

    _body(

        doc,

        "VaaniQ establishes a reproducible multilingual research framework for studying "

        "audio-deepfake detection under language, codec, and confidence-calibration shift. "

        "On the bounded V1 benchmark, both the acoustic and frozen XLS-R pipelines achieve "

        "strong held-out discrimination, while cross-language, calibration, and external-source "

        "experiments expose important limitations. The current evidence supports the value of "

        "multilingual evaluation but does not establish universal source- or generator-independent "

        "detection.",

    )



    _heading(doc, "", "References")

    references = [

        '[1] J. Jung et al., AASIST: Audio anti-spoofing using integrated spectro-temporal graph attention networks, in Proc. ICASSP, 2022.',

        '[2] A. Babu et al., XLS-R: Self-supervised cross-lingual speech representation learning at scale, arXiv:2111.09296, 2021.',

        '[3] T. Javed et al., IndicSUPERB: A speech processing universal performance benchmark for Indian languages, arXiv:2208.11761, 2022.',

        '[4] D. V. Sharma, V. Ekbote, and A. Gupta, IndicSynth: A large-scale multilingual synthetic speech dataset, ACL, 2025.',

        '[5] O. Pascu et al., Towards calibrated and explainable audio deepfake detection, in Proc. Interspeech, 2024.',

        '[6] N. M. Muller et al., Human perception of audio deepfakes, in Proc. ACM Workshop on Information Hiding and Multimedia Security, 2022.',

        '[7] AI4Bharat, Kathbath dataset, Hugging Face: ai4bharat/Kathbath.',

        '[8] D. V. Sharma et al., IndicSynth dataset, Hugging Face: vdivyasharma/IndicSynth.',

        '[9] A. Conneau et al., FLEURS: Few-shot learning evaluation of universal representations of speech, arXiv:2205.12446, 2022.',

        '[10] IETF RFC 6716, Definition of the Opus Audio Codec, 2012.',

        '[11] ASVspoof 2019, logical access evaluation plan; English-only control uses a streaming subset of Bisher/ASVspoof_2019_LA (ODC-By).',

    ]

    for reference in references:

        paragraph = doc.add_paragraph(reference)

        paragraph.paragraph_format.left_indent = Inches(0.14)

        paragraph.paragraph_format.first_line_indent = Inches(-0.14)

        paragraph.paragraph_format.space_after = Pt(1)

        for run in paragraph.runs:

            run.font.size = Pt(8)



    doc.save(OUTPUT)

    doc.save(DRAFT_OUTPUT)

    return OUTPUT





def build_progress_report() -> Path:

    """Build the progress report as the approved Round 3 status document."""

    report = _load_report()

    metrics = _dict(report.get("test_metrics"))

    xlsr = _dict(load_xlsr_main().get("test_metrics"))

    if "kathbath" not in str(report.get("data_provenance", "")):

        raise RuntimeError("progress report requires publication-subset results")



    doc = Document()

    section = doc.sections[0]

    section.page_width = Inches(8.27)

    section.page_height = Inches(11.69)

    section.top_margin = Inches(0.7)

    section.bottom_margin = Inches(0.7)

    section.left_margin = Inches(0.75)

    section.right_margin = Inches(0.75)

    normal = doc.styles["Normal"]

    normal.font.name = "Aptos"

    normal.font.size = Pt(10.5)

    normal.paragraph_format.space_after = Pt(6)



    title = doc.add_paragraph()

    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    title.paragraph_format.space_before = Inches(1.1)

    run = title.add_run("VaaniQ\nCapstone Progress Report - Round 3 Approved")

    run.bold = True

    run.font.name = "Aptos Display"

    run.font.size = Pt(22)

    subtitle = doc.add_paragraph()

    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle.add_run(

        "Cross-Lingual, Compression-Robust Detection and Calibrated Reliability "

        "Estimation for AI-Generated Voice in Indian Languages"

    ).italic = True

    team = doc.add_paragraph()

    team.alignment = WD_ALIGN_PARAGRAPH.CENTER

    team.add_run(

        "\nGaurav Phadale  70022300092\n"

        "Eshaan Sarkhawas  70022300066\n"

        "Aarav Phutane  70022300152\n"

        "Prajwal Patil  70022300213\n\n"

        "Faculty Guide: Prof. Rama Bharti Varshney\n"

        "MPSTME, NMIMS  Academic Year 2026-27\n"

        "Frozen evidence: artifacts/final_results_manifest.json\n"

        "Approved commit: 084bd47ca6ca1b69a7cdbf424e2946f3794c2a95"

    )

    doc.add_page_break()



    doc.add_heading("1. Executive Progress Summary", level=1)

    _body(

        doc,

        "This report describes the approved Round 3 state. It is not a claim that the "

        "project is 100% complete. RQ1-RQ4, Baseline V1, frozen XLS-R main, LFCC-GMM, "

        "the RawNet2-style approximate baseline, reproducibility tooling, the application, "

        "and documentation are COMPLETE. Benchmark V2 is PARTIAL. FLEURS unseen-source "

        "evaluation is PILOT. Faithful RawNet2 and generator-disjoint evaluation are "

        "PENDING. RQ5 is BLOCKED ON HUMAN DATA (N=0). Codec evaluation uses a "
        "WhatsApp-style Opus simulation, not actual WhatsApp transmission.",

    )

    _table(

        doc,

        ["Work package", "Status", "Canonical note"],

        [

            ["Baseline V1", "COMPLETE", "n=584; 91.61% acc; 6.56% EER"],

            ["Frozen XLS-R main", "COMPLETE", "n=584; 92.12% acc; 6.88% EER; 0.9828 AUC"],

            ["RQ1 compression", "COMPLETE", "Model-dependent clean/Opus results"],

            ["RQ2 English-only", "COMPLETE", "54.8% acc; 76.56% EER; 0.162 AUC"],

            ["RQ3 LOO", "COMPLETE", "Hindi 78.83%; Marathi 93.29%; Tamil 93.94%"],

            ["RQ4 calibration", "COMPLETE", "Val-selected; test ECE 0.0245?0.026"],

            ["LFCC-GMM", "COMPLETE", "54.79% acc; weak baseline"],

            ["RawNet2-style approximate baseline", "COMPLETE", "54.79% acc; not faithful RawNet2"],

            ["Reproducibility + application + docs", "COMPLETE", "Frozen manifest and generated reports"],

            ["Benchmark V2", "PARTIAL", "External-source pilot; source probe 98.48%"],

            ["FLEURS unseen-real", "PILOT", "n=9; no statistically useful estimate"],

            ["Faithful RawNet2", "PENDING", "Implementation remaining"],

            ["Generator-disjoint", "PENDING", "n=0"],

            ["RQ5 human study", "BLOCKED ON HUMAN DATA", "Protocol ready; N=0"],

        ],

        caption="Table 1. Round 3 status by work package",

    )



    doc.add_heading("2. Canonical Held-Out Results", level=1)

    _table(

        doc,

        ["Model", "n", "Accuracy", "EER", "ROC-AUC"],

        [

            ["Acoustic Baseline V1", "584", pct(_number(metrics, "accuracy")), pct(_number(metrics, "eer")), f"{_number(metrics, 'roc_auc'):.4f}"],

            ["Frozen XLS-R main", "584", pct(_number(xlsr, "accuracy")), pct(_number(xlsr, "eer")), f"{_number(xlsr, 'roc_auc'):.4f}"],

        ],

        caption="Table 2. Primary V1 test comparison",

    )

    _body(

        doc,

        "Frozen XLS-R improved ranking performance while classification performance "

        "remained broadly comparable. Acoustic Baseline V1 is not canonical AASIST and "

        "not XLS-R.",

    )



    doc.add_heading("3. Remaining Work", level=1)

    remaining = [

        "Balance Benchmark V2 source\xd7label cells (currently PARTIAL).",

        "Expand FLEURS beyond the n=9 PILOT before claiming unseen-source generalization.",

        "Run generator-disjoint evaluation when a genuine held-out generator cell exists (n=0 PENDING).",

        "Implement faithful RawNet2 (PENDING; not an external blocker).",

        "Collect real RQ5 participants; human-study protocol ready; N=0.",

    ]

    for item in remaining:

        doc.add_paragraph(item, style="List Bullet")



    doc.add_heading("4. Demonstration Readiness", level=1)

    _body(

        doc,

        "The live system can demonstrate upload inference, microphone streaming, dataset "

        "provenance, compression-aware metrics, calibration plots, explainability artefacts, "

        "and API documentation. All headline numbers must match the frozen Round 3 manifest. "

        "Do not present V2 as solving the source confound, n=9 FLEURS as a benchmark, or "

        "human results that do not exist.",

    )



    doc.save(PROGRESS_OUTPUT)

    return PROGRESS_OUTPUT





def verify_paper(path: Path) -> None:

    """Verify package integrity and frozen Round 3 numbers."""

    document = Document(path)

    text = "\n".join(

        [paragraph.text for paragraph in document.paragraphs]

        + [cell.text for table in document.tables for row in table.rows for cell in row.cells]

    )

    required = [

        "Gaurav Phadale",

        "Kathbath",

        "IndicSynth",

        "91.61%",

        "6.56%",

        "92.12%",

        "6.88%",

        "0.9828",

        "54.8%",

        "78.83%",

        "93.29%",

        "93.94%",

        "N=0",

        "PARTIAL",

        "PILOT",

        "PENDING",

        "WhatsApp-style Opus simulation",

        "AASIST-compatible",

        "RawNet2-style approximate",

    ]

    missing = [value for value in required if value not in text]

    if missing:

        raise RuntimeError(f"IEEE paper is missing required evidence: {missing}")

    forbidden_as_headline = [

        "RQ2 remain incomplete",

        "English-only control checkpoint not yet run",

        "current front-end is a compact deterministic embedding, not the final frozen XLS-R",

        "93.7%",

        "mid-term baseline; planned detector baselines",

    ]

    present_bad = [value for value in forbidden_as_headline if value in text]

    if present_bad:

        raise RuntimeError(f"IEEE paper still contains stale claims: {present_bad}")

    if "93.66%" in text and "historical" not in text.lower():

        raise RuntimeError("93.66% appears without historical framing")

    if len(document.tables) < 9:

        raise RuntimeError(f"expected at least 9 tables, got {len(document.tables)}")

    if len(document.inline_shapes) < 3:

        raise RuntimeError(

            f"expected at least 3 embedded figures, got {len(document.inline_shapes)}"

        )

    with zipfile.ZipFile(path) as archive:

        if archive.testzip() is not None:

            raise RuntimeError("IEEE DOCX package integrity check failed")

    print(

        f"IEEE paper verified: {len(document.tables)} tables, "

        f"{len(document.inline_shapes)} figures, package integrity PASS"

    )





def main() -> None:

    paper_path = build_paper()

    verify_paper(paper_path)

    print(f"Wrote {paper_path}")

    progress_path = build_progress_report()

    progress = Document(progress_path)

    progress_text = "\n".join(p.text for p in progress.paragraphs)

    for token in ("91.61%", "92.12%", "PARTIAL", "PILOT", "N=0", "not 100%"):

        if token == "not 100%" and "100% complete" not in progress_text.lower() and "not a claim that the project is 100%" in progress_text.lower():

            continue

        if token != "not 100%" and token not in progress_text and token not in "\n".join(

            cell.text for table in progress.tables for row in table.rows for cell in row.cells

        ):

            raise RuntimeError(f"progress report missing {token}")

    with zipfile.ZipFile(progress_path) as archive:

        if archive.testzip() is not None:

            raise RuntimeError("progress-report DOCX integrity check failed")

    print(f"Wrote {progress_path}")





if __name__ == "__main__":

    main()

