#!/usr/bin/env python3
"""Cross-check master DOCX text against train_report.json and live API."""

from __future__ import annotations

import json
import re
import sys
import urllib.request
import zipfile
from pathlib import Path

from docx import Document
from PIL import Image, ImageChops

REPO = Path(__file__).resolve().parents[1]
TRAIN_REPORT = REPO / "models" / "checkpoints" / "xlsr_aasist" / "train_report.json"
DOCX_CANDIDATES = [
    REPO / "docs" / "VaaniQ_Master_Presentation_FINAL.docx",
    REPO / "docs" / "VaaniQ_Master_Presentation_COMPLETE.docx",
    REPO / "docs" / "VaaniQ_Master_Presentation.docx",
]
API = "http://127.0.0.1:8001"


def load_report() -> dict[str, object]:
    with TRAIN_REPORT.open(encoding="utf-8") as fh:
        raw = json.load(fh)
    return raw if isinstance(raw, dict) else {}


def fetch_json(url: str) -> dict[str, object]:
    with urllib.request.urlopen(url, timeout=30) as resp:
        raw = json.loads(resp.read().decode())
    return raw if isinstance(raw, dict) else {}


def resolve_docx() -> Path:
    existing = [path for path in DOCX_CANDIDATES if path.is_file()]
    if not existing:
        raise FileNotFoundError("No master DOCX found")
    return max(existing, key=lambda path: path.stat().st_mtime)


def doc_text(doc: Document) -> str:
    chunks: list[str] = []
    for paragraph in doc.paragraphs:
        chunks.append(paragraph.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                chunks.append(cell.text)
    return "\n".join(chunks)


def pct(value: float) -> str:
    return f"{value:.1%}"


def verify_figure_margins(paths: list[Path], *, minimum_pixels: int = 4) -> list[str]:
    """Return figure names whose visible content touches an image boundary."""
    failures: list[str] = []
    for path in paths:
        with Image.open(path) as source:
            image = source.convert("RGB")
        background = Image.new("RGB", image.size, "white")
        difference = ImageChops.difference(image, background).convert("L")
        foreground = difference.point(lambda value: 255 if value > 10 else 0)
        bounds = foreground.getbbox()
        if bounds is None:
            failures.append(f"{path.name}: blank")
            continue
        left, top, right, bottom = bounds
        width, height = image.size
        if (
            left < minimum_pixels
            or top < minimum_pixels
            or right > width - minimum_pixels
            or bottom > height - minimum_pixels
        ):
            failures.append(
                f"{path.name}: content bounds={bounds}, image={image.size}"
            )
    return failures


def main() -> int:
    report = load_report()
    test = report.get("test_metrics")
    test = test if isinstance(test, dict) else {}
    eval_metrics = report.get("eval_metrics")
    if not isinstance(eval_metrics, dict):
        eval_metrics = test

    docx_path = resolve_docx()
    doc = Document(str(docx_path))
    text = doc_text(doc)

    publication_subset = "kathbath" in str(report.get("data_provenance", ""))
    required_strings = [
        pct(float(report.get("test_accuracy", 0))),
        f"{float(report.get('test_ece', 0)):.3f}",
        f"{float(report.get('test_brier', 0)):.3f}",
        str(report.get("n_clips", 900)),
        str(report.get("n_train", 540)),
        str(report.get("n_val", 180)),
        str(report.get("n_test", 180)),
        "Hindi",
        "Marathi",
        "Tamil",
        "Kathbath" if publication_subset else "synthetic",
        "IndicSynth" if publication_subset else "demonstration",
        "Speaker-disjoint" if publication_subset else "Data provenance",
    ]
    missing = [item for item in required_strings if item not in text]
    if missing:
        print("FAIL: DOCX missing expected metric strings:")
        for item in missing:
            print(f"  - {item}")
        return 1

    matrix = eval_metrics.get("confusion_matrix")
    if isinstance(matrix, list):
        for row in matrix:
            if isinstance(row, list):
                for cell in row:
                    if str(int(cell)) not in text:
                        print(f"WARN: confusion cell {cell} not found in DOCX text")

    per_lang = eval_metrics.get("per_language")
    if isinstance(per_lang, dict):
        for lang in ("hi", "mr", "ta"):
            block = per_lang.get(lang)
            if isinstance(block, dict):
                acc = float(block.get("accuracy", 0))
                if pct(acc) not in text and f"{acc:.1%}" not in text:
                    print(f"FAIL: per-language accuracy for {lang} ({pct(acc)}) not in DOCX")
                    return 1

    inline = len(doc.inline_shapes)
    tables = len(doc.tables)
    if inline < 24:
        print(f"FAIL: expected >=24 inline images, got {inline}")
        return 1
    if tables < 10:
        print(f"FAIL: expected >=10 tables, got {tables}")
        return 1

    with zipfile.ZipFile(docx_path) as archive:
        media = [name for name in archive.namelist() if name.startswith("word/media/")]
        if archive.testzip() is not None:
            print("FAIL: DOCX zip integrity check failed")
            return 1
        if len(media) < 24:
            print(f"FAIL: expected >=24 media files, got {len(media)}")
            return 1

    verified_figs = list((REPO / "docs" / "assets" / "verified_figures").glob("*.png"))
    if len(verified_figs) < 9:
        print(f"FAIL: expected >=9 verified figure PNGs, got {len(verified_figs)}")
        return 1
    expected_figures = {
        "performance_overview.png",
        "language_breakdown.png",
        "condition_breakdown.png",
        "confusion_matrix.png",
        "calibration_errors.png",
        "calibration_pre_post.png",
        "reliability_diagram.png",
        "system_architecture.png",
        "ml_pipeline.png",
    }
    missing_figures = sorted(
        expected_figures - {path.name for path in verified_figs}
    )
    if missing_figures:
        print(f"FAIL: verified figures missing: {missing_figures}")
        return 1
    margin_failures = verify_figure_margins(
        [path for path in verified_figs if path.name in expected_figures]
    )
    if margin_failures:
        print("FAIL: cropped or blank verified figures:")
        for failure in margin_failures:
            print(f"  - {failure}")
        return 1

    try:
        pipeline = fetch_json(f"{API}/api/v1/metrics/pipeline")
        api_acc = float(pipeline.get("test_accuracy", -1))
        report_acc = float(report.get("test_accuracy", -2))
        if abs(api_acc - report_acc) > 1e-4:
            print(f"FAIL: API test_accuracy {api_acc} != report {report_acc}")
            return 1
        health = fetch_json(f"{API}/api/v1/admin/status")
        if str(health.get("status", "")).lower() not in {"ok", "healthy", "ready"}:
            print(f"WARN: API admin status is {health.get('status')}")
    except OSError as exc:
        print(f"WARN: API not reachable for live cross-check: {exc}")

    if re.search(r"\b100\.0%\b.*conference", text, re.IGNORECASE):
        print("WARN: possible overclaim of 100% conference result")

    print("PASS: master DOCX audit")
    print(f"  file: {docx_path.name}")
    print(f"  inline_images: {inline}")
    print(f"  tables: {tables}")
    print(f"  media_files: {len(media)}")
    print(f"  verified_figures: {len(verified_figs)}")
    print("  figure_margin_audit: PASS")
    print(f"  test_accuracy: {pct(float(report.get('test_accuracy', 0)))}")
    print(f"  test_eer: {pct(float(report.get('test_eer', 0)))}")
    print(f"  test_ece: {float(report.get('test_ece', 0)):.3f}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
